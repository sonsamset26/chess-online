# -*- coding: utf-8 -*-
"""
Script tự động sinh tài liệu Word: SRS_Document_Chess_Online.docx
Dự án: Nền tảng Đánh Cờ Vua Trực tuyến Thời gian thực tích hợp Trí tuệ Nhân tạo
Tác giả: Phan Hồng Sơn - MSV: 174765 - Lớp: 65PM-CNVLVH
Quy cách: Hướng dẫn IEEE 830, tham khảo mẫu tài liệu thực tế và quy định HUCE
"""

import os
import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
import docx
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top="D3D3D3", bottom="D3D3D3", left="D3D3D3", right="D3D3D3"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for b_name, b_color in borders.items():
        if b_color:
            b_element = OxmlElement(f'w:{b_name}')
            b_element.set(qn('w:val'), 'single')
            b_element.set(qn('w:sz'), '4')
            b_element.set(qn('w:space'), '0')
            b_element.set(qn('w:color'), b_color)
            tcBorders.append(b_element)
        else:
            b_element = OxmlElement(f'w:{b_name}')
            b_element.set(qn('w:val'), 'none')
            tcBorders.append(b_element)
    tcPr.append(tcBorders)

def build_srs_document():
    doc = docx.Document()

    # Cấu hình lề trang theo quy định: Trên 20mm, Dưới 20mm, Trái 30mm, Phải 20mm
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)
        section.page_width = Mm(210)
        section.page_height = Mm(297)

    # Cấu hình kiểu chữ mặc định: Times New Roman 13pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    def add_p(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, font_size=13, space_before=0, space_after=4, color_rgb=(0, 0, 0)):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = RGBColor(*color_rgb)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 32, 96) # Navy Blue
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13.5)
        run.bold = True
        run.font.color.rgb = RGBColor(15, 60, 120) # Blue
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor(30, 80, 150)
        return p

    def add_code_block(title, code_str):
        # Tiêu đề khối mã
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(6)
        p_t.paragraph_format.space_after = Pt(2)
        run_t = p_t.add_run(f"Mã định nghĩa ({title}):")
        run_t.bold = True
        run_t.italic = True
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = RGBColor(50, 50, 50)

        # Hộp chứa mã dạng bảng 1 ô
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(6.3)
        set_cell_background(cell, "F8F9FA")
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        set_cell_borders(cell, top="D0D7DE", bottom="D0D7DE", left="D0D7DE", right="D0D7DE")

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        run = p.add_run(code_str.strip())
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(36, 41, 47)

        add_p("", space_before=2, space_after=4)

    # ==============================================================================
    # 1. TRANG BÌA CHÍNH (THEO MẪU BÌA ĐỒ ÁN / TTTN HUCE)
    # ==============================================================================
    add_p("TRƯỜNG ĐẠI HỌC XÂY DỰNG HÀ NỘI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=14, space_before=15, space_after=2)
    add_p("KHOA CÔNG NGHỆ THÔNG TIN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=13, space_after=2)
    add_p("BỘ MÔN CÔNG NGHỆ PHẦN MỀM", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=13, space_after=15)
    add_p("-----------------------------------------", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    add_p("TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=18, space_before=20, space_after=10, color_rgb=(0, 32, 96))
    add_p("HỆ THỐNG CỜ VUA TRỰC TUYẾN THỜI GIAN THỰC TÍCH HỢP TRÍ TUỆ NHÂN TẠO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=15, space_after=10, color_rgb=(15, 60, 120))
    add_p("NỀN TẢNG THI ĐẤU, HỌC TẬP VÀ GIẢI ĐỐ CỜ VUA TRÊN ĐÁM MÂY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, italic=True, font_size=12.5, space_after=60)

    # Bảng thông tin sinh viên và đơn vị hướng dẫn
    table_cover = doc.add_table(rows=7, cols=2)
    table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cover.autofit = False

    cover_data = [
        ("Sinh viên thực hiện:", "Phan Hồng Sơn"),
        ("Mã số sinh viên:", "174765"),
        ("Lớp / Khóa học:", "65PM-CNVLVH (Khóa 65)"),
        ("Đơn vị hướng dẫn:", "Công ty Cổ phần VTI"),
        ("Cán bộ hướng dẫn tại đơn vị:", "Đinh Văn Đông (Trưởng nhóm Kỹ thuật)"),
        ("Giảng viên hướng dẫn:", "ThS. Nguyễn Hải Dương"),
        ("Phiên bản tài liệu:", "1.0 (Ban hành tháng 09/2026)")
    ]

    for idx, (label, val) in enumerate(cover_data):
        row = table_cover.rows[idx]
        p_lbl = row.cells[0].paragraphs[0]
        p_lbl.add_run(label).bold = True
        p_lbl.runs[0].font.name = 'Times New Roman'
        p_lbl.runs[0].font.size = Pt(12.5)
        row.cells[0].width = Inches(2.6)

        p_val = row.cells[1].paragraphs[0]
        p_val.add_run(val)
        p_val.runs[0].font.name = 'Times New Roman'
        p_val.runs[0].font.size = Pt(12.5)
        p_val.runs[0].bold = (idx in [0, 4, 5])
        row.cells[1].width = Inches(3.7)

    add_p("", space_before=40)
    add_p("HÀ NỘI – NĂM 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=13, space_before=50)

    doc.add_page_break()

    # ==============================================================================
    # 2. MỤC LỤC VÀ DANH MỤC
    # ==============================================================================
    add_h1("MỤC LỤC TỔNG QUAN")
    
    toc_items = [
        ("1. GIỚI THIỆU", "Trang 3"),
        ("   1.1 Mục đích tài liệu", "Trang 3"),
        ("   1.2 Phạm vi hệ thống", "Trang 3"),
        ("   1.3 Bảng thuật ngữ và định nghĩa", "Trang 4"),
        ("   1.4 Tài liệu tham khảo", "Trang 4"),
        ("   1.5 Tổng quan cấu trúc tài liệu", "Trang 4"),
        ("2. YÊU CẦU CHỨC NĂNG", "Trang 5"),
        ("   2.1 Các tác nhân hệ thống", "Trang 5"),
        ("   2.2 Danh mục chức năng tổng quan", "Trang 5"),
        ("   2.3 Biểu đồ Use Case tổng quan", "Trang 7"),
        ("   2.4 Biểu đồ Use Case phân rã cho từng tác nhân", "Trang 9"),
        ("   2.5 Biểu đồ trình tự các luồng nghiệp vụ cốt lõi", "Trang 12"),
        ("   2.6 Bảng đặc tả Use Case chi tiết", "Trang 17"),
        ("   2.7 Bảng ma trận kịch bản kiểm thử", "Trang 26"),
        ("3. YÊU CẦU PHI CHỨC NĂNG", "Trang 28"),
        ("   3.1 Giao diện người dùng", "Trang 28"),
        ("   3.2 Hiệu năng hệ thống", "Trang 28"),
        ("   3.3 Độ tin cậy và tính sẵn sàng", "Trang 28"),
        ("   3.4 An toàn và bảo mật dữ liệu", "Trang 29"),
        ("   3.5 Khả năng mở rộng", "Trang 29"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        run1 = p.add_run(item)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        if item.startswith("1.") or item.startswith("2.") or item.startswith("3."):
            run1.bold = True
            run1.font.color.rgb = RGBColor(0, 32, 96)
        
        # Dấu chấm chấm nối sang trang
        dots_len = max(5, 75 - len(item))
        run_dots = p.add_run(" " + "." * dots_len + " ")
        run_dots.font.name = 'Times New Roman'
        run_dots.font.size = Pt(10)
        run_dots.font.color.rgb = RGBColor(120, 120, 120)

        run2 = p.add_run(page)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11.5)
        run2.bold = True

    doc.add_page_break()

    # ==============================================================================
    # 3. CHƯƠNG 1: GIỚI THIỆU
    # ==============================================================================
    add_h1("1. GIỚI THIỆU")

    add_h2("1.1 Mục đích")
    add_p("Tài liệu đặc tả yêu cầu phần mềm này trình bày chi tiết các yêu cầu chức năng, yêu cầu phi chức năng cùng các mô hình hành vi, kịch bản tương tác và tiêu chí kiểm thử cho hệ thống Nền tảng Đánh Cờ Vua Trực tuyến Thời gian thực. Tài liệu được xây dựng nhằm làm cơ sở kỹ thuật thống nhất giữa tác giả, giảng viên hướng dẫn và đơn vị thực tập trong các giai đoạn phân tích thiết kế, lập trình, kiểm thử và nghiệm thu sản phẩm phần mềm.")

    add_h2("1.2 Phạm vi hệ thống")
    add_p("Hệ thống là một giải pháp trực tuyến đa người dùng trên nền tảng Web, phục vụ nhu cầu thi đấu cờ vua, rèn luyện kỹ năng và giải trí tương tác cao:")
    add_p("• Thi đấu với máy tính: Tích hợp thuật toán Negamax kết hợp cắt tỉa Alpha-Beta chạy trên luồng ngầm của trình duyệt, cung cấp 3 cấp độ chơi từ dễ đến khó mà không làm gián đoạn giao diện bàn cờ.")
    add_p("• Thi đấu trực tuyến giữa người với người:")
    add_p("  - Chế độ ghép trận xếp hạng tự động: Tìm kiếm đối thủ có trình độ tương đương trong hàng chờ máy chủ và tính toán biến thiên điểm Elo sau mỗi ván đấu theo luật của Liên đoàn Cờ vua Quốc tế.")
    add_p("  - Chế độ tạo phòng bạn bè: Tạo phòng đấu riêng với mã phòng gồm 6 ký tự để người chơi giao hữu mà không ảnh hưởng đến điểm xếp hạng.")
    add_p("• Đồng bộ thời gian thực và kiểm soát luật cờ:")
    add_p("  - Toàn bộ nước đi được máy chủ kiểm tra tính hợp lệ trước khi chấp nhận và phát tán tới đối thủ qua kết nối thời gian thực.")
    add_p("  - Giải thuật đồng hồ thi đấu hướng sự kiện: Máy chủ chỉ ghi nhận các mốc thời gian chuyển lượt, tính toán và trừ thời gian suy nghĩ một lần khi nhận nước đi hợp lệ, kết hợp bộ hẹn giờ giám sát hết giờ nhằm tiết kiệm tài nguyên vi xử lý máy chủ.")
    add_p("  - Cơ chế khôi phục trạng thái ván cờ trong 45 giây: Khi người chơi mất kết nối mạng hoặc tải lại trang web, máy chủ duy trì ván cờ và đếm lùi thời gian cho phép người chơi vào lại mà không bị xử thua ngay lập tức.")
    add_p("• Tổ chức giải đấu loại trực tiếp: Hỗ trợ quy mô 4 hoặc 8 kỳ thủ với sơ đồ phân nhánh tự động. Khi ván đấu chính hòa, hệ thống tự động khởi tạo ván phụ thi đấu nhanh với màu quân đảo chiều và áp dụng lợi thế hòa cờ cho bên cầm quân Đen để xác định người chiến thắng đi tiếp.")
    add_p("• Xem lại và phân tích thế cờ: Tích hợp Stockfish engine phiên bản WebAssembly giúp người chơi xem lại lịch sử nước đi, đánh giá chất lượng từng nước cờ thông qua chỉ số tổn thất ưu thế và phát hiện các sai lầm chiến thuật.")
    add_p("• Học tập và giải đố cờ vua: Ngân hàng 30 bài tập cờ thế chiến thuật phân bổ theo các dải trình độ từ 1000 đến 2050 Elo cùng hệ thống bài giảng tương tác giúp người mới bắt đầu làm quen luật chơi và nâng cao trình độ.")
    add_p("• Hồ sơ kỳ thủ và phân tích phong cách bằng học máy: Trích xuất 8 chỉ số đặc trưng thi đấu từ lịch sử các ván cờ (độ mất điểm thế cờ trung bình và tỉ lệ sai sót qua 3 giai đoạn khai cuộc, trung cuộc, tàn cuộc; thời gian suy nghĩ trung bình; tỉ lệ sai sót khi cạn giờ). Áp dụng mô hình K-Means và bộ chuẩn hóa StandardScaler để phân cụm người chơi vào 4 nhóm phong cách: Tiến công, Toàn diện, Đột biến và Phòng thủ. Hệ thống đồng thời chẩn đoán giai đoạn thi đấu có hiệu suất thấp nhất để đề xuất các bài tập cờ thế chiến thuật phù hợp và hiển thị trực quan qua biểu đồ mạng nhện SVG 8 trục.")
    add_p("• Quản lý tài khoản và bảo mật: Xác thực danh tính với cơ chế bảo mật kết hợp mã truy cập ngắn hạn lưu trong bộ nhớ tạm và mã làm mới dài hạn lưu trong cookie bảo mật chỉ đọc, ngăn ngừa các nguy cơ khai thác trái phép qua mạng.")

    add_h2("1.3 Bảng thuật ngữ và định nghĩa")
    
    table_glossary = doc.add_table(rows=1, cols=2)
    table_glossary.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_glossary.autofit = False

    hdr_g = table_glossary.rows[0].cells
    hdr_g[0].paragraphs[0].add_run("Thuật ngữ").bold = True
    hdr_g[1].paragraphs[0].add_run("Ý nghĩa và định nghĩa nghiệp vụ").bold = True
    hdr_g[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_g[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for c in hdr_g:
        set_cell_background(c, "1F4E79")
        set_cell_margins(c, 120, 120, 160, 160)
        set_cell_borders(c, "0F2537", "0F2537", "0F2537", "0F2537")

    hdr_g[0].width = Inches(1.8)
    hdr_g[1].width = Inches(4.5)

    glossary_data = [
        ("FIDE", "Liên đoàn Cờ vua Quốc tế, cơ quan ban hành luật cờ và quy tắc tính điểm xếp hạng toàn cầu."),
        ("Elo", "Hệ thống đánh giá trình độ kỳ thủ dựa trên kết quả thi đấu đối kháng, điểm số tăng khi thắng và giảm khi thua."),
        ("FEN", "Ký hiệu quy định cấu trúc chuỗi ký tự thể hiện vị trí toàn bộ quân cờ và trạng thái bàn cờ tại một thời điểm."),
        ("PGN", "Định dạng văn bản quy định chuỗi nước đi và thông tin đi kèm của một ván cờ vua hoàn chỉnh."),
        ("UCI", "Giao thức truyền thông giữa giao diện người dùng và các công cụ tính toán nước cờ."),
        ("WASM", "Định dạng mã nhị phân hiệu năng cao thực thi trực tiếp trên trình duyệt web."),
        ("Stockfish engine", "Động cơ phân tích thế cờ mã nguồn mở mã hóa sang WebAssembly thực thi trên Web Worker của trình duyệt."),
        ("CPL", "Độ mất mát ưu thế của nước cờ đo bằng một phần trăm giá trị quy đổi của quân Tốt."),
        ("PvAI", "Chế độ thi đấu đối kháng giữa người chơi và máy tính."),
        ("PvP", "Chế độ thi đấu đối kháng trực tiếp giữa hai người chơi thực."),
        ("WebSocket", "Giao thức mạng cho phép truyền thông hai chiều thời gian thực giữa máy chủ và máy khách qua một kết nối duy nhất."),
        ("JWT", "Định dạng mã định danh bảo mật dùng để xác thực và ủy quyền truy cập trong ứng dụng web."),
        ("Ván phụ Armageddon", "Hình thức thi đấu ván phụ nhanh để phân định thắng thua khi hòa cờ, bên Trắng có nhiều thời gian hơn nhưng bên Đen có ưu thế hòa là thắng."),
        ("K-Means", "Thuật toán học máy không giám sát dùng để nhóm tập dữ liệu thành K cụm dựa trên khoảng cách hình học tới tâm cụm."),
        ("StandardScaler", "Kỹ thuật chuẩn hóa dữ liệu đưa kỳ vọng về 0 và phương sai về 1, loại bỏ sai lệch thang đo giữa các biến đặc trưng."),
        ("Radar Chart", "Biểu đồ mạng nhện đa giác biểu diễn đồng thời nhiều biến định lượng trên cùng một hệ tọa độ cực.")
    ]

    for idx, (term, desc) in enumerate(glossary_data):
        row = table_glossary.add_row().cells
        row[0].paragraphs[0].add_run(term).bold = True
        row[1].paragraphs[0].add_run(desc)
        bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
        for c in row:
            set_cell_background(c, bg)
            set_cell_margins(c, 90, 90, 130, 130)
            set_cell_borders(c, "D3D3D3", "D3D3D3", "D3D3D3", "D3D3D3")
            c.paragraphs[0].runs[0].font.name = 'Times New Roman'
            c.paragraphs[0].runs[0].font.size = Pt(11.5)
        row[0].width = Inches(1.8)
        row[1].width = Inches(4.5)

    add_p("", space_before=4)

    add_h2("1.4 Tài liệu tham khảo")
    add_p("1. IEEE Std 830-1998: Hướng dẫn thực hành khuyến nghị cho đặc tả yêu cầu phần mềm của Viện Kỹ sư Điện và Điện tử.")
    add_p("2. FIDE Handbook: Luật thi đấu cờ vua của Liên đoàn Cờ vua Quốc tế.")
    add_p("3. Tài liệu kiến trúc và hướng dẫn đặc tả yêu cầu phần mềm tham khảo thực tế.")
    add_p("4. Báo cáo khảo sát nghiệp vụ và tài liệu kỹ thuật dự án Nền tảng Đánh Cờ Vua Trực tuyến.")

    add_h2("1.5 Tổng quan tài liệu")
    add_p("Tài liệu được cấu trúc thành 3 phần chính:")
    add_p("• Phần 1 - Giới thiệu: Trình bày mục đích, phạm vi, bảng thuật ngữ, tài liệu tham khảo và tổng quan cấu trúc.")
    add_p("• Phần 2 - Yêu cầu chức năng: Mô tả chi tiết các tác nhân hệ thống, danh mục chức năng, biểu đồ Use Case tổng quan và phân rã, biểu đồ trình tự 3 luồng nghiệp vụ cốt lõi, bảng đặc tả chi tiết 7 Use Case chính và ma trận 10 kịch bản kiểm thử thực tế.")
    add_p("• Phần 3 - Yêu cầu phi chức năng: Xác định các chỉ tiêu kỹ thuật về giao diện người dùng, hiệu năng, độ tin cậy và tính sẵn sàng, an toàn bảo mật dữ liệu cùng khả năng mở rộng hệ thống.")

    doc.add_page_break()

    # ==============================================================================
    # 4. CHƯƠNG 2: YÊU CẦU CHỨC NĂNG
    # ==============================================================================
    add_h1("2. YÊU CẦU CHỨC NĂNG")

    add_h2("2.1 Các tác nhân hệ thống")
    add_p("Hệ thống bao gồm 3 nhóm tác nhân chính:")
    add_p("1. Khách: Người dùng truy cập website mà chưa thực hiện đăng nhập tài khoản. Khách có thể trải nghiệm các tính năng mở như đấu cờ với máy tính, giải các bài tập cờ thế cơ bản, học các bài giảng nhập môn và xem bảng xếp hạng kỳ thủ.")
    add_p("2. Người chơi có tài khoản: Người dùng đã hoàn tất đăng ký và đăng nhập vào hệ thống. Tác nhân này được sử dụng toàn bộ tính năng của khách, đồng thời có thể tham gia hàng chờ ghép trận xếp hạng có tính điểm Elo, tạo hoặc tham gia phòng đấu bạn bè bằng mã phòng, ghi danh thi đấu giải đấu loại trực tiếp, xem lại và phân tích chi tiết ván cờ đã đấu, cũng như lưu trữ tiến trình học tập cá nhân.")
    add_p("3. Quản trị viên: Người dùng có thẩm quyền cao nhất trong hệ thống, chịu trách nhiệm quản lý tài khoản người chơi, khóa hoặc mở khóa tài khoản vi phạm, quản lý ngân hàng bài học và thế cờ chiến thuật, đồng thời giám sát các phòng đấu và giải đấu đang diễn ra.")

    add_h2("2.2 Danh mục chức năng tổng quan")
    
    table_fn = doc.add_table(rows=1, cols=4)
    table_fn.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_fn.autofit = False

    hdr_fn = table_fn.rows[0].cells
    hdr_fn[0].paragraphs[0].add_run("Nhóm chức năng").bold = True
    hdr_fn[1].paragraphs[0].add_run("Mã").bold = True
    hdr_fn[2].paragraphs[0].add_run("Tên chức năng").bold = True
    hdr_fn[3].paragraphs[0].add_run("Tác nhân thực hiện").bold = True
    for c in hdr_fn:
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "1F4E79")
        set_cell_margins(c, 100, 100, 140, 140)
        set_cell_borders(c, "0F2537", "0F2537", "0F2537", "0F2537")

    hdr_fn[0].width = Inches(1.8)
    hdr_fn[1].width = Inches(0.8)
    hdr_fn[2].width = Inches(2.3)
    hdr_fn[3].width = Inches(1.4)

    fn_list = [
        ("Quản lý tài khoản", "FN-01", "Đăng ký tài khoản mới", "Khách"),
        ("Quản lý tài khoản", "FN-02", "Đăng nhập hệ thống", "Khách, Người chơi"),
        ("Quản lý tài khoản", "FN-03", "Đăng xuất", "Người chơi, Quản trị viên"),
        ("Quản lý tài khoản", "FN-04", "Xem và cập nhật hồ sơ cá nhân", "Người chơi"),
        ("Quản lý tài khoản", "FN-05", "Xem bảng xếp hạng Elo", "Tất cả"),
        ("Đấu với máy tính", "FN-06", "Lựa chọn cấp độ chơi và màu quân", "Khách, Người chơi"),
        ("Đấu với máy tính", "FN-07", "Tương tác bàn cờ và nhận nước đi từ máy", "Khách, Người chơi"),
        ("Đấu trực tuyến", "FN-08", "Ghép trận ngẫu nhiên tính điểm Elo", "Người chơi"),
        ("Đấu trực tuyến", "FN-09", "Tạo phòng thi đấu bạn bè và nhận mã phòng", "Người chơi"),
        ("Đấu trực tuyến", "FN-10", "Nhập mã phòng để tham gia phòng bạn bè", "Người chơi"),
        ("Đấu trực tuyến", "FN-11", "Kiểm soát luật cờ và đồng bộ nước đi thời gian thực", "Hệ thống"),
        ("Đấu trực tuyến", "FN-12", "Vận hành đồng hồ thi đấu hướng sự kiện", "Hệ thống"),
        ("Đấu trực tuyến", "FN-13", "Tạm giữ ván đấu và khôi phục trong 45 giây", "Hệ thống, Người chơi"),
        ("Giải đấu", "FN-14", "Tạo giải đấu loại trực tiếp 4 hoặc 8 người", "Người chơi"),
        ("Giải đấu", "FN-15", "Tham gia giải đấu bằng mã mời", "Người chơi"),
        ("Giải đấu", "FN-16", "Tự động sinh nhánh đấu và điều phối ván cờ", "Hệ thống"),
        ("Giải đấu", "FN-17", "Tự động kích hoạt ván phụ khi có kết quả hòa", "Hệ thống, Người chơi"),
        ("Giải đấu", "FN-18", "Chuyển vòng đấu sau thời gian đếm ngược 30 giây", "Hệ thống, Người chơi"),
        ("Phân tích ván cờ", "FN-19", "Lưu trữ lịch sử ván cờ và chuỗi nước đi", "Hệ thống"),
        ("Phân tích ván cờ", "FN-20", "Xem lại biên bản ván đấu từng nước đi", "Người chơi, Khách"),
        ("Phân tích ván cờ", "FN-21", "Đánh giá chất lượng nước cờ bằng công cụ phân tích", "Người chơi"),
        ("Học cờ và thế cờ", "FN-22", "Xem danh sách bài học và thực hành tương tác", "Khách, Người chơi"),
        ("Học cờ và thế cờ", "FN-23", "Giải bài tập thế cờ chiến thuật", "Khách, Người chơi"),
        ("Quản trị hệ thống", "FN-24", "Quản lý tài khoản người dùng và khóa tài khoản", "Quản trị viên"),
        ("Quản trị hệ thống", "FN-25", "Quản lý ngân hàng bài học và bài tập cờ thế", "Quản trị viên"),
        ("Quản trị hệ thống", "FN-26", "Giám sát trạng thái phòng đấu và giải đấu", "Quản trị viên")
    ]

    for idx, (grp, code, name, actor) in enumerate(fn_list):
        row = table_fn.add_row().cells
        row[0].paragraphs[0].add_run(grp)
        row[1].paragraphs[0].add_run(code).bold = True
        row[2].paragraphs[0].add_run(name)
        row[3].paragraphs[0].add_run(actor)
        bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
        for c in row:
            set_cell_background(c, bg)
            set_cell_margins(c, 70, 70, 110, 110)
            set_cell_borders(c, "D3D3D3", "D3D3D3", "D3D3D3", "D3D3D3")
            c.paragraphs[0].runs[0].font.name = 'Times New Roman'
            c.paragraphs[0].runs[0].font.size = Pt(11)
        row[0].width = Inches(1.8)
        row[1].width = Inches(0.8)
        row[2].width = Inches(2.3)
        row[3].width = Inches(1.4)

    add_p("", space_before=6)

    # 2.3 Biểu đồ Use Case tổng quan
    add_h2("2.3 Biểu đồ Use Case tổng quan")
    add_p("Biểu đồ Use Case tổng quan thể hiện bức tranh toàn cảnh về tương tác giữa 3 tác nhân (Khách, Người chơi có tài khoản, Quản trị viên) với các cụm chức năng chính của hệ thống:")

    plantuml_overall = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam shadowing false
skinparam defaultFontName Arial

actor "Khách" as Guest
actor "Người chơi có tài khoản" as User
actor "Quản trị viên" as Admin

rectangle "Hệ thống Chess Online" {
  usecase "Đăng ký tài khoản" as UC_Register
  usecase "Đăng nhập hệ thống" as UC_Login
  usecase "Xem bảng xếp hạng" as UC_Leaderboard
  usecase "Đấu với máy tính" as UC_PlayAI
  usecase "Giải bài tập cờ thế" as UC_Puzzles
  usecase "Học cờ qua bài giảng" as UC_Lessons

  usecase "Cập nhật hồ sơ cá nhân" as UC_Profile
  usecase "Ghép trận đấu xếp hạng" as UC_Matchmaking
  usecase "Tạo phòng bạn bè" as UC_CreateRoom
  usecase "Tham gia phòng bạn bè" as UC_JoinRoom
  usecase "Tham gia giải đấu" as UC_Tournament
  usecase "Xem lại và phân tích ván cờ" as UC_Analyze

  usecase "Quản lý người chơi" as UC_ManageUsers
  usecase "Quản lý ngân hàng bài học & thế cờ" as UC_ManageContent
  usecase "Giám sát phòng đấu & giải đấu" as UC_Monitor
}

Guest <|-- User

Guest --> UC_Register
Guest --> UC_Login
Guest --> UC_Leaderboard
Guest --> UC_PlayAI
Guest --> UC_Puzzles
Guest --> UC_Lessons

User --> UC_Profile
User --> UC_Matchmaking
User --> UC_CreateRoom
User --> UC_JoinRoom
User --> UC_Tournament
User --> UC_Analyze

Admin --> UC_Login
Admin --> UC_ManageUsers
Admin --> UC_ManageContent
Admin --> UC_Monitor
@enduml"""
    add_code_block("PlantUML - Biểu đồ Use Case tổng quan", plantuml_overall)

    mermaid_overall = """flowchart LR
    subgraph Users ["Tác nhân hệ thống"]
        Guest["Khách"]
        User["Người chơi có tài khoản"]
        Admin["Quản trị viên"]
    end

    subgraph System ["Hệ thống Chess Online"]
        UC01["Đăng ký & Đăng nhập"]
        UC02["Xem bảng xếp hạng"]
        UC03["Đấu với máy tính"]
        UC04["Giải bài tập thế cờ"]
        UC05["Học cờ qua bài giảng"]
        UC06["Cập nhật hồ sơ cá nhân"]
        UC07["Ghép trận đấu xếp hạng"]
        UC08["Tạo & tham gia phòng bạn bè"]
        UC09["Tham gia giải đấu loại trực tiếp"]
        UC10["Xem lại & phân tích ván cờ"]
        UC11["Quản lý người dùng"]
        UC12["Quản lý bài học & thế cờ"]
        UC13["Giám sát phòng & giải đấu"]
    end

    Guest --> UC01
    Guest --> UC02
    Guest --> UC03
    Guest --> UC04
    Guest --> UC05

    User --> UC01
    User --> UC02
    User --> UC03
    User --> UC04
    User --> UC05
    User --> UC06
    User --> UC07
    User --> UC08
    User --> UC09
    User --> UC10

    Admin --> UC01
    Admin --> UC11
    Admin --> UC12
    Admin --> UC13"""
    add_code_block("Mermaid - Biểu đồ Use Case tổng quan", mermaid_overall)

    # 2.4 Biểu đồ Use Case phân rã
    add_h2("2.4 Biểu đồ Use Case phân rã cho từng tác nhân")

    add_h3("2.4.1 Phân rã Use Case cho tác nhân Khách")
    add_p("Tác nhân Khách được tiếp cận các chức năng cơ bản nhằm làm quen giao diện và trải nghiệm cờ vua trước khi đăng ký:")

    plantuml_guest = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam shadowing false
skinparam defaultFontName Arial

actor "Khách" as Guest

rectangle "Phân hệ Người dùng vãng lai" {
  usecase "Đăng ký tài khoản" as UC_Reg
  usecase "Đăng nhập tài khoản" as UC_Log
  usecase "Xem bảng xếp hạng Elo" as UC_Rank
  
  usecase "Đấu với máy tính" as UC_AI
  usecase "Chọn cấp độ dễ, vừa, khó" as UC_Level
  usecase "Chọn màu quân Trắng, Đen, Ngẫu nhiên" as UC_Color
  
  usecase "Giải bài tập cờ thế" as UC_Puz
  usecase "Xem bài tập gợi ý" as UC_PuzHint
  
  usecase "Học bài học cờ vua" as UC_Les
  usecase "Tương tác bàn cờ bài học" as UC_LesBoard
}

Guest --> UC_Reg
Guest --> UC_Log
Guest --> UC_Rank
Guest --> UC_AI
Guest --> UC_Puz
Guest --> UC_Les

UC_AI ..> UC_Level : <<include>>
UC_AI ..> UC_Color : <<include>>
UC_Puz ..> UC_PuzHint : <<extend>>
UC_Les ..> UC_LesBoard : <<include>>
@enduml"""
    add_code_block("PlantUML - Phân rã tác nhân Khách", plantuml_guest)

    mermaid_guest = """flowchart LR
    Guest["Khách"]

    subgraph GuestScope ["Chức năng dành cho Khách"]
        UC_Reg["Đăng ký tài khoản"]
        UC_Log["Đăng nhập hệ thống"]
        UC_Rank["Xem bảng xếp hạng Elo"]
        
        UC_AI["Đấu với máy tính"]
        UC_Level["Chọn cấp độ chơi"]
        UC_Color["Chọn màu quân"]
        
        UC_Puz["Giải bài tập cờ thế"]
        UC_Hint["Xem gợi ý thế cờ"]
        
        UC_Les["Học cờ cơ bản"]
        UC_LesStep["Thực hành nước đi theo bài"]
    end

    Guest --> UC_Reg
    Guest --> UC_Log
    Guest --> UC_Rank
    Guest --> UC_AI
    Guest --> UC_Puz
    Guest --> UC_Les

    UC_AI -.->|include| UC_Level
    UC_AI -.->|include| UC_Color
    UC_Puz -.->|extend| UC_Hint
    UC_Les -.->|include| UC_LesStep"""
    add_code_block("Mermaid - Phân rã tác nhân Khách", mermaid_guest)

    add_h3("2.4.2 Phân rã Use Case cho tác nhân Người chơi có tài khoản")
    add_p("Người chơi có tài khoản là đối tượng chính sử dụng hệ thống, sở hữu đầy đủ quyền hạn tham gia thi đấu đối kháng trực tuyến, giải đấu và phân tích thế cờ:")

    plantuml_player = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam shadowing false
skinparam defaultFontName Arial

actor "Người chơi có tài khoản" as Player

rectangle "Phân hệ Người chơi Đã đăng ký" {
  usecase "Quản lý thông tin cá nhân" as UC_Prof
  usecase "Xem lịch sử ván cờ" as UC_Hist
  
  usecase "Ghép trận đấu xếp hạng" as UC_Match
  usecase "Hủy hàng chờ ghép trận" as UC_CancelQueue
  usecase "Khôi phục ván cờ khi mất mạng" as UC_Reconnect
  
  usecase "Tạo phòng bạn bè" as UC_Create
  usecase "Tham gia phòng bằng mã" as UC_Join
  
  usecase "Tham gia giải đấu loại trực tiếp" as UC_Tour
  usecase "Thi đấu ván chính trong giải" as UC_TourMatch
  usecase "Thi đấu ván phụ khi hòa" as UC_TieBreak
  usecase "Chờ đếm ngược chuyển vòng" as UC_Countdown
  
  usecase "Xem lại và phân tích ván cờ" as UC_Review
  usecase "Đánh giá chất lượng nước cờ" as UC_Eval
}

Player --> UC_Prof
Player --> UC_Hist
Player --> UC_Match
Player --> UC_Create
Player --> UC_Join
Player --> UC_Tour
Player --> UC_Review

UC_Match ..> UC_CancelQueue : <<extend>>
UC_Match ..> UC_Reconnect : <<extend>>
UC_Tour ..> UC_TourMatch : <<include>>
UC_Tour ..> UC_TieBreak : <<extend>>
UC_Tour ..> UC_Countdown : <<include>>
UC_Review ..> UC_Eval : <<extend>>
@enduml"""
    add_code_block("PlantUML - Phân rã Người chơi có tài khoản", plantuml_player)

    mermaid_player = """flowchart LR
    Player["Người chơi có tài khoản"]

    subgraph PlayerScope ["Chức năng Người chơi có tài khoản"]
        UC_Prof["Quản lý hồ sơ & Thống kê"]
        UC_Hist["Xem danh sách ván đấu đã chơi"]
        
        UC_Match["Ghép trận đấu xếp hạng"]
        UC_Cancel["Hủy tìm trận"]
        UC_Recon["Khôi phục ván cờ 45 giây"]
        
        UC_Room["Đấu phòng bạn bè"]
        UC_MakeRoom["Tạo mã phòng 6 ký tự"]
        UC_EnterRoom["Nhập mã vào phòng"]
        
        UC_Tour["Tham gia giải đấu loại trực tiếp"]
        UC_Round["Thi đấu vòng loại trực tiếp"]
        UC_Arma["Thi đấu ván phụ khi hòa"]
        UC_NextRound["Chuyển vòng sau 30 giây"]
        
        UC_Rev["Xem lại và phân tích ván cờ"]
        UC_Stockfish["Đánh giá chất lượng nước đi"]
    end

    Player --> UC_Prof
    Player --> UC_Hist
    Player --> UC_Match
    Player --> UC_Room
    Player --> UC_Tour
    Player --> UC_Rev

    UC_Match -.->|extend| UC_Cancel
    UC_Match -.->|extend| UC_Recon
    UC_Room -.->|include| UC_MakeRoom
    UC_Room -.->|include| UC_EnterRoom
    UC_Tour -.->|include| UC_Round
    UC_Tour -.->|extend| UC_Arma
    UC_Tour -.->|include| UC_NextRound
    UC_Rev -.->|extend| UC_Stockfish"""
    add_code_block("Mermaid - Phân rã Người chơi có tài khoản", mermaid_player)

    add_h3("2.4.3 Phân rã Use Case cho tác nhân Quản trị viên")
    add_p("Quản trị viên thực hiện quản lý người dùng, duyệt nội dung và giám sát hoạt động thời gian thực của máy chủ:")

    plantuml_admin = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam shadowing false
skinparam defaultFontName Arial

actor "Quản trị viên" as Admin

rectangle "Phân hệ Quản trị Hệ thống" {
  usecase "Đăng nhập quyền quản trị" as UC_AdminLogin
  usecase "Quản lý người chơi" as UC_UserMgmt
  usecase "Tìm kiếm người dùng" as UC_FindUser
  usecase "Khóa tài khoản vi phạm" as UC_LockUser
  usecase "Mở khóa tài khoản" as UC_UnlockUser
  
  usecase "Quản lý ngân hàng thế cờ" as UC_PuzMgmt
  usecase "Thêm bài tập thế cờ mới" as UC_AddPuz
  usecase "Kiểm tra chuỗi FEN hợp lệ" as UC_ValidFEN
  
  usecase "Quản lý bài giảng cờ vua" as UC_LesMgmt
  usecase "Thêm và chỉnh sửa bài giảng" as UC_EditLes
  
  usecase "Giám sát giải đấu & phòng đấu" as UC_Monitor
  usecase "Xem danh sách phòng đang chơi" as UC_LiveRooms
}

Admin --> UC_AdminLogin
Admin --> UC_UserMgmt
Admin --> UC_PuzMgmt
Admin --> UC_LesMgmt
Admin --> UC_Monitor

UC_UserMgmt ..> UC_FindUser : <<include>>
UC_UserMgmt ..> UC_LockUser : <<extend>>
UC_UserMgmt ..> UC_UnlockUser : <<extend>>

UC_PuzMgmt ..> UC_AddPuz : <<include>>
UC_AddPuz ..> UC_ValidFEN : <<include>>

UC_LesMgmt ..> UC_EditLes : <<include>>
UC_Monitor ..> UC_LiveRooms : <<include>>
@enduml"""
    add_code_block("PlantUML - Phân rã Quản trị viên", plantuml_admin)

    mermaid_admin = """flowchart LR
    Admin["Quản trị viên"]

    subgraph AdminScope ["Chức năng Quản trị viên"]
        UC_AdmLog["Đăng nhập quản trị"]
        
        UC_Users["Quản lý người dùng"]
        UC_Search["Tìm kiếm người dùng"]
        UC_Lock["Khóa tài khoản vi phạm"]
        UC_Unlock["Mở khóa tài khoản"]
        
        UC_Puzzles["Quản lý bài tập thế cờ"]
        UC_AddPuz["Thêm bài tập thế cờ mới"]
        UC_ValFEN["Kiểm tra tính hợp lệ chuỗi FEN"]
        
        UC_Lessons["Quản lý bài giảng cờ vua"]
        UC_EditLes["Biên tập nội dung bài giảng"]
        
        UC_Mon["Giám sát hệ thống"]
        UC_ViewRooms["Xem danh sách phòng đấu trực tiếp"]
    end

    Admin --> UC_AdmLog
    Admin --> UC_Users
    Admin --> UC_Puzzles
    Admin --> UC_Lessons
    Admin --> UC_Mon

    UC_Users -.->|include| UC_Search
    UC_Users -.->|extend| UC_Lock
    UC_Users -.->|extend| UC_Unlock

    UC_Puzzles -.->|include| UC_AddPuz
    UC_AddPuz -.->|include| UC_ValFEN

    UC_Lessons -.->|include| UC_EditLes
    UC_Mon -.->|include| UC_ViewRooms"""
    add_code_block("Mermaid - Phân rã Quản trị viên", mermaid_admin)

    doc.add_page_break()

    # 2.5 Biểu đồ trình tự các luồng nghiệp vụ cốt lõi
    add_h2("2.5 Biểu đồ trình tự các luồng nghiệp vụ cốt lõi")

    add_h3("2.5.1 Luồng 1: Ghép trận xếp hạng và đồng bộ nước đi thời gian thực")
    add_p("Mô tả nghiệp vụ: Hai người chơi gửi yêu cầu tìm trận xếp hạng. Cổng kết nối đưa vào hàng chờ, so khớp mức điểm Elo, gán ngẫu nhiên màu quân và tạo phòng đấu. Trong ván cờ, mỗi nước đi gửi lên máy chủ đều trải qua quy trình kiểm tra danh tính kết nối, tính hợp lệ của lượt đi, tính toán và trừ thời gian suy nghĩ vào quỹ giờ dựa trên mốc thời gian máy chủ, đồng thời xác thực luật cờ trước khi phát tán tới đối thủ. Khi kết thúc ván đấu, hệ thống tính toán biến thiên điểm Elo và cập nhật cơ sở dữ liệu.")

    plantuml_seq1 = """@startuml
skinparam shadowing false
skinparam defaultFontName Arial
autonumber

actor "Người chơi A" as PlayerA
actor "Người chơi B" as PlayerB
participant "Cổng kết nối\\n(WebSocket)" as Gateway
participant "Bộ điều phối trận đấu\\n(MatchService)" as MatchService
participant "Kiểm tra luật cờ\\n(ChessValidator)" as Validator
database "Cơ sở dữ liệu\\n(MongoDB)" as Database

PlayerA -> Gateway: Gửi yêu cầu tìm trận xếp hạng (find_match)
activate Gateway
Gateway -> MatchService: Đưa vào hàng chờ (waitingQueue)
activate MatchService
MatchService --> Gateway: Xác nhận đang tìm đối thủ
Gateway --> PlayerA: Hiển thị trạng thái đang tìm trận

PlayerB -> Gateway: Gửi yêu cầu tìm trận xếp hạng (find_match)
Gateway -> MatchService: Đưa vào hàng chờ (waitingQueue)
MatchService -> MatchService: So khớp điểm Elo giữa A và B
MatchService -> MatchService: Tạo mã phòng và chỉ định bên cầm Trắng/Đen
MatchService -> Database: Tạo bản ghi ván đấu mới (status = PLAYING)
activate Database
Database --> MatchService: Bản ghi tạo thành công
deactivate Database

MatchService --> Gateway: Thông báo ghép trận thành công
deactivate MatchService

Gateway -> PlayerA: Phát sự kiện match_found (màu Trắng, thời gian 10 phút)
Gateway -> PlayerB: Phát sự kiện match_found (màu Đen, thời gian 10 phút)
Gateway -> Gateway: Khởi tạo đồng hồ máy chủ (turnStartedAt = now)

== Quá trình thực hiện và đồng bộ nước đi ==

PlayerA -> Gateway: Gửi nước đi (make_move: e2 -> e4)
Gateway -> Gateway: Xác thực định danh kết nối và lượt đi của bên Trắng
Gateway -> Gateway: Tính thời gian suy nghĩ = now - turnStartedAt
Gateway -> Gateway: Trừ thời gian suy nghĩ vào quỹ giờ bên Trắng
Gateway -> Validator: Kiểm tra tính hợp lệ của nước đi e2-e4
activate Validator
Validator --> Gateway: Nước đi hợp lệ, trả về chuỗi FEN mới
deactivate Validator

Gateway -> Gateway: Cập nhật turnStartedAt mới cho lượt bên Đen
Gateway -> Gateway: Thiết lập bộ hẹn giờ giám sát hết giờ cho bên Đen

Gateway -> PlayerA: Phát sự kiện receive_move (FEN mới, thời gian hai bên)
Gateway -> PlayerB: Phát sự kiện receive_move (FEN mới, thời gian hai bên)

== Kết thúc ván cờ và cập nhật xếp hạng ==

PlayerB -> Gateway: Gửi nước đi dẫn đến thế Chiếu hết (Checkmate)
Gateway -> Validator: Xác nhận tình trạng ván đấu (isCheckmate = true)
activate Validator
Validator --> Gateway: Xác nhận bên Đen chiến thắng
deactivate Validator

Gateway -> MatchService: Kết thúc ván đấu và tính toán biến thiên điểm Elo
activate MatchService
MatchService -> Database: Cập nhật kết quả ván cờ, FEN cuối cùng và chuỗi PGN
activate Database
MatchService -> Database: Cập nhật điểm Elo mới cho Người chơi A và Người chơi B
Database --> MatchService: Cập nhật thành công
deactivate Database
MatchService --> Gateway: Trả về kết quả tổng kết ván cờ
deactivate MatchService

Gateway -> PlayerA: Phát sự kiện game_over (kết quả thua, điểm Elo mới)
Gateway -> PlayerB: Phát sự kiện game_over (kết quả thắng, điểm Elo mới)
deactivate Gateway
@enduml"""
    add_code_block("PlantUML - Trình tự Luồng 1: Ghép trận và đồng bộ nước đi", plantuml_seq1)

    mermaid_seq1 = """sequenceDiagram
    autonumber
    actor PA as Người chơi A
    actor PB as Người chơi B
    participant GW as Cổng kết nối WebSocket
    participant MS as Bộ điều phối trận đấu
    participant CV as Kiểm tra luật cờ
    participant DB as Cơ sở dữ liệu MongoDB

    PA->>GW: Gửi yêu cầu tìm trận xếp hạng
    GW->>MS: Thêm Người chơi A vào hàng chờ
    MS-->>GW: Xác nhận vào hàng chờ
    GW-->>PA: Hiển thị giao diện đang tìm đối thủ

    PB->>GW: Gửi yêu cầu tìm trận xếp hạng
    GW->>MS: Thêm Người chơi B vào hàng chờ
    MS->>MS: So khớp điểm Elo tương thích
    MS->>MS: Khởi tạo phòng đấu, gán A cầm Trắng, B cầm Đen
    MS->>DB: Tạo bản ghi ván đấu mới
    DB-->>MS: Ghi nhận thành công
    MS-->>GW: Thông báo trận đấu sẵn sàng

    GW->>PA: Phát match_found (cầm Trắng, 10 phút)
    GW->>PB: Phát match_found (cầm Đen, 10 phút)
    GW->>GW: Ghi nhận mốc thời gian máy chủ ban đầu

    Note over PA,GW: Đồng bộ nước đi thời gian thực
    PA->>GW: Gửi nước đi e2-e4
    GW->>GW: Xác thực lượt đi & tính thời gian đã suy nghĩ
    GW->>GW: Trừ thời gian vào quỹ giờ bên Trắng
    GW->>CV: Kiểm tra tính hợp lệ của nước đi
    CV-->>GW: Nước đi hợp lệ, trả về FEN mới
    GW->>GW: Đổi lượt cho Đen, cập nhật mốc thời gian mới
    GW->>PA: Phát receive_move (FEN mới, quỹ thời gian)
    GW->>PB: Phát receive_move (FEN mới, quỹ thời gian)

    Note over PA,DB: Kết thúc ván cờ
    PB->>GW: Gửi nước đi tạo thế chiếu hết
    GW->>CV: Xác nhận tình trạng chiếu hết
    CV-->>GW: Kết luận Đen thắng do chiếu hết
    GW->>MS: Tính toán biến thiên điểm Elo hai bên
    MS->>DB: Lưu biên bản PGN, kết quả và điểm Elo mới
    DB-->>MS: Lưu thành công
    MS-->>GW: Hoàn tất xử lý
    GW->>PA: Phát game_over (Thua, điểm Elo mới)
    GW->>PB: Phát game_over (Thắng, điểm Elo mới)"""
    add_code_block("Mermaid - Trình tự Luồng 1: Ghép trận và đồng bộ nước đi", mermaid_seq1)

    add_h3("2.5.2 Luồng 2: Tạm mất kết nối mạng và khôi phục ván cờ trong 45 giây")
    add_p("Mô tả nghiệp vụ: Khi một bên gặp sự cố mạng hoặc vô tình tải lại trang web, máy chủ không xử thua ngay mà chuyển phòng sang trạng thái chờ kết nối lại, đồng thời kích hoạt bộ đếm thời gian ân hạn 45 giây và gửi cảnh báo kèm đồng hồ đếm lùi tới đối thủ. Khi người chơi kết nối lại trong 45 giây, hệ thống đồng bộ lại toàn bộ dữ liệu bàn cờ và tiếp tục ván đấu. Nếu quá 45 giây, người chơi mới bị xử thua do bỏ cuộc.")

    plantuml_seq2 = """@startuml
skinparam shadowing false
skinparam defaultFontName Arial
autonumber

actor "Người chơi A\\n(Gặp sự cố)" as PlayerA
actor "Người chơi B\\n(Đang chờ)" as PlayerB
participant "Cổng kết nối\\n(WebSocket)" as Gateway
participant "Bộ đếm ân hạn 45s\\n(GracePeriodTimer)" as Timer
participant "Bộ nhớ ván đấu\\n(ActiveRooms)" as RoomState
database "Cơ sở dữ liệu\\n(MongoDB)" as Database

Note over PlayerA, PlayerB: Ván cờ đang diễn ra bình thường tại phòng đấu
PlayerA -[#red]x Gateway: Đứt kết nối mạng / Tải lại trình duyệt (disconnect)
activate Gateway
Gateway -> RoomState: Kiểm tra trạng thái phòng đấu hiện tại
activate RoomState
RoomState --> Gateway: Ván cờ đang ở trạng thái PLAYING
deactivate RoomState

Gateway -> RoomState: Chuyển trạng thái sang RECONNECTING
Gateway -> Timer: Kích hoạt bộ đếm thời gian ân hạn 45 giây
activate Timer
Gateway -> PlayerB: Phát sự kiện player_disconnected\\n(Đối thủ rớt mạng, đếm lùi 45s)

alt Trường hợp 1: Người chơi A kết nối lại thành công trong 45 giây
  PlayerA -> Gateway: Kết nối lại WebSocket thành công
  PlayerA -> Gateway: Gửi yêu cầu reconnect_match (kèm userId, roomId)
  Gateway -> RoomState: Xác thực quyền tham gia của Người chơi A
  activate RoomState
  RoomState --> Gateway: Xác nhận hợp lệ
  deactivate RoomState

  Gateway -> Timer: Hủy bỏ bộ đếm thời gian ân hạn 45 giây
  deactivate Timer
  Gateway -> RoomState: Cập nhật mã kết nối mới cho Người chơi A
  Gateway -> RoomState: Chuyển trạng thái phòng trở lại PLAYING

  Gateway -> PlayerA: Phát sự kiện sync_game_state\\n(FEN, lượt đi, quỹ thời gian hai bên, PGN)
  Gateway -> PlayerB: Phát sự kiện player_reconnected\\n(Đối thủ đã vào lại, tiếp tục thi đấu)

else Trường hợp 2: Quá thời hạn 45 giây không kết nối lại
  Timer -> Gateway: Kích hoạt sự kiện timeout hết 45 giây
  activate Timer
  deactivate Timer
  Gateway -> RoomState: Đóng phòng đấu do người chơi bỏ cuộc
  Gateway -> Database: Ghi nhận Người chơi A thua cuộc do bỏ kết nối
  activate Database
  Gateway -> Database: Cập nhật điểm Elo thắng cho Người chơi B
  Database --> Gateway: Lưu thành công
  deactivate Database

  Gateway -> PlayerB: Phát sự kiện game_over (Thắng cuộc do đối thủ bỏ trận)
end
deactivate Gateway
@enduml"""
    add_code_block("PlantUML - Trình tự Luồng 2: Khôi phục ván cờ 45 giây", plantuml_seq2)

    mermaid_seq2 = """sequenceDiagram
    autonumber
    actor PA as Người chơi A (Rớt mạng)
    actor PB as Người chơi B (Đang chờ)
    participant GW as Cổng kết nối WebSocket
    participant TM as Bộ đếm ân hạn 45 giây
    participant RS as Bộ nhớ ván đấu ActiveRooms
    participant DB as Cơ sở dữ liệu MongoDB

    Note over PA,PB: Ván đấu đang trong trạng thái thi đấu bình thường
    PA-xGW: Mất kết nối mạng hoặc tải lại trang web
    GW->>RS: Kiểm tra trạng thái phòng đấu
    RS-->>GW: Trạng thái hiện tại đang thi đấu
    GW->>RS: Chuyển trạng thái phòng sang RECONNECTING
    GW->>TM: Khởi động bộ đếm lùi 45 giây
    GW->>PB: Phát cảnh báo đối thủ mất kết nối kèm đồng hồ 45s

    alt Trường hợp 1: Người chơi A vào lại kịp trước 45 giây
        PA->>GW: Mở lại web, gửi yêu cầu reconnect_match
        GW->>RS: Xác thực danh tính và phòng đấu
        RS-->>GW: Xác thực hợp lệ
        GW->>TM: Hủy bộ đếm thời gian ân hạn
        GW->>RS: Đổi trạng thái phòng về PLAYING, gán kết nối mới
        GW->>PA: Phát sync_game_state (Hoàn trả ván cờ, FEN, quỹ giờ)
        GW->>PB: Phát player_reconnected (Đối thủ đã vào lại)
    else Trường hợp 2: Hết 45 giây không vào lại
        TM->>GW: Báo sự kiện hết thời gian ân hạn 45 giây
        GW->>RS: Hủy phòng đấu
        GW->>DB: Ghi nhận A thua do bỏ cuộc, cộng điểm Elo cho B
        DB-->>GW: Lưu kết quả thành công
        GW->>PB: Phát game_over (B thắng do A bỏ cuộc)
    end"""
    add_code_block("Mermaid - Trình tự Luồng 2: Khôi phục ván cờ 45 giây", mermaid_seq2)

    add_h3("2.5.3 Luồng 3: Vòng đời giải đấu loại trực tiếp")
    add_p("Mô tả nghiệp vụ: Chủ phòng tạo giải đấu 4 người và nhận mã mời. Các kỳ thủ tham gia đầy đủ, chủ phòng bấm bắt đầu giải đấu. Hệ thống tự động tạo 2 trận đấu vòng bán kết. Trận 1 có người thắng tiến vào chung kết. Trận 2 có kết quả hòa cờ: hệ thống tự động kích hoạt ván phụ thi đấu nhanh với màu quân đảo chiều (Trắng 5 phút, Đen 4 phút, Đen có ưu thế hòa là thắng) để xác định người đi tiếp. Sau khi toàn bộ các trận vòng bán kết kết thúc, hệ thống kích hoạt đồng hồ đếm ngược 30 giây nghỉ chuẩn bị trước khi tự động khởi tạo trận chung kết.")

    plantuml_seq3 = """@startuml
skinparam shadowing false
skinparam defaultFontName Arial
autonumber

actor "Chủ phòng giải đấu" as Host
actor "Các kỳ thủ tham gia" as Players
participant "Cổng kết nối\\n(WebSocket)" as Gateway
participant "Dịch vụ giải đấu\\n(TournamentService)" as TourService
participant "Phòng đấu ván cờ\\n(MatchGateway)" as MatchGW
database "Cơ sở dữ liệu\\n(MongoDB)" as Database

Host -> Gateway: Gửi yêu cầu tạo giải đấu (size = 4)
activate Gateway
Gateway -> TourService: Khởi tạo Tournament mới
activate TourService
TourService -> Database: Lưu giải đấu mới (status = WAITING, mã mời 6 ký tự)
activate Database
Database --> TourService: Lưu thành công
deactivate Database
TourService --> Gateway: Trả về mã mời tham gia
Gateway --> Host: Hiển thị phòng chờ giải đấu và mã mời

Players -> Gateway: Gửi yêu cầu tham gia (join_tournament: mã mời)
Gateway -> TourService: Thêm kỳ thủ vào danh sách tham gia
TourService --> Gateway: Danh sách cập nhật (4/4 người)
Gateway -> Host: Phát sự kiện tournament_updated
Gateway -> Players: Phát sự kiện tournament_updated

Host -> Gateway: Phát lệnh bắt đầu giải đấu (start_tournament)
Gateway -> TourService: Sinh sơ đồ nhánh đấu loại trực tiếp (Round 1: 2 trận)
TourService -> Database: Cập nhật trạng thái IN_PROGRESS
activate Database
Database --> TourService: Lưu thành công
deactivate Database
TourService --> Gateway: Trả về danh sách cặp đấu vòng 1
Gateway -> Host: Phát sự kiện tournament_started kèm thông tin nhánh đấu
Gateway -> Players: Phát sự kiện tournament_started kèm thông tin nhánh đấu

== Tiến hành Vòng 1: Trận 1 có người thắng, Trận 2 hòa cờ ==

TourService -> MatchGW: Tự động khởi tạo Trận 1 và Trận 2
activate MatchGW
MatchGW --> Gateway: Hai phòng đấu vòng bán kết sẵn sàng
Gateway -> Players: Mời các cặp đấu vào bàn cờ tương ứng

MatchGW -> TourService: Trận 1 kết thúc (Người chơi 1 thắng Người chơi 2)
TourService -> Database: Cập nhật người thắng Trận 1 vào nhánh Chung kết

MatchGW -> TourService: Trận 2 kết thúc với kết quả Hòa cờ (winnerColor = draw)
TourService -> TourService: Phát hiện giải đấu loại trực tiếp cần phân định thắng thua
TourService -> MatchGW: Tự động kích hoạt Ván phụ thi đấu nhanh
MatchGW -> MatchGW: Đảo màu quân (Trắng thành Đen, Đen thành Trắng)
MatchGW -> MatchGW: Thiết lập thời gian lệch (Trắng 5 phút, Đen 4 phút)
MatchGW -> MatchGW: Áp dụng quy tắc bên Đen hưởng ưu thế hòa là thắng

Gateway -> Players: Thông báo ván chính hòa, bắt đầu ván phụ phân định
MatchGW -> TourService: Ván phụ kết thúc với kết quả hòa -> Xử Đen thắng
TourService -> Database: Cập nhật người thắng Trận 2 vào nhánh Chung kết
deactivate MatchGW

== Giai đoạn đếm ngược 30 giây và chuyển sang Vòng Chung kết ==

TourService -> TourService: Kiểm tra toàn bộ các trận Vòng 1 đã hoàn tất
TourService --> Gateway: Yêu cầu khởi động thời gian chuẩn bị chuyển vòng
Gateway -> Players: Phát sự kiện round_countdown (Đếm ngược 30 giây nghỉ)

... Sau thời gian đếm ngược 30 giây hoàn tất ...

Gateway -> TourService: Hết 30 giây chuẩn bị, yêu cầu tạo vòng kế tiếp
TourService -> MatchGW: Tự động tạo phòng thi đấu Trận Chung kết
activate MatchGW
MatchGW --> Gateway: Phòng chung kết sẵn sàng
Gateway -> Players: Mời 2 kỳ thủ xuất sắc vào thi đấu Chung kết

MatchGW -> TourService: Trận Chung kết kết thúc (Xác định Nhà vô địch)
deactivate MatchGW
TourService -> Database: Cập nhật championId và chuyển status = FINISHED
activate Database
Database --> TourService: Lưu hoàn tất
deactivate Database
TourService --> Gateway: Thông báo kết thúc giải đấu
deactivate TourService

Gateway -> Host: Phát sự kiện tournament_finished (Công bố Nhà vô địch)
Gateway -> Players: Phát sự kiện tournament_finished (Công bố Nhà vô địch)
deactivate Gateway
@enduml"""
    add_code_block("PlantUML - Trình tự Luồng 3: Vòng đời giải đấu loại trực tiếp", plantuml_seq3)

    mermaid_seq3 = """sequenceDiagram
    autonumber
    actor Host as Chủ phòng
    actor Players as Các kỳ thủ tham gia
    participant GW as Cổng kết nối WebSocket
    participant TS as Dịch vụ giải đấu
    participant MG as Bộ điều phối trận đấu
    participant DB as Cơ sở dữ liệu MongoDB

    Host->>GW: Tạo giải đấu quy mô 4 người
    GW->>TS: Khởi tạo giải đấu mới
    TS->>DB: Lưu trạng thái WAITING, sinh mã mời 6 số
    DB-->>TS: Lưu thành công
    TS-->>GW: Trả về mã mời
    GW-->>Host: Hiển thị phòng chờ và mã mời

    Players->>GW: Nhập mã tham gia giải đấu
    GW->>TS: Ghi nhận người chơi (Đủ 4/4 người)
    TS-->>GW: Cập nhật danh sách phòng chờ
    GW->>Host: Báo đủ người, sẵn sàng bắt đầu

    Host->>GW: Ra lệnh bắt đầu giải đấu
    GW->>TS: Khởi tạo nhánh đấu Bán kết (2 trận)
    TS->>DB: Đổi trạng thái giải sang IN_PROGRESS
    TS-->>GW: Danh sách 2 cặp đấu vòng 1
    GW->>Players: Mời các kỳ thủ vào bàn đấu

    Note over Players,MG: Diễn biến Vòng Bán kết
    TS->>MG: Tạo phòng đấu Trận 1 và Trận 2
    MG->>TS: Trận 1 kết thúc có người thắng rõ ràng
    TS->>DB: Đưa người thắng Trận 1 vào trận Chung kết

    MG->>TS: Trận 2 kết thúc với kết quả Hòa cờ
    Note over TS,MG: Kích hoạt ván phụ phân định thắng thua
    TS->>MG: Khởi tạo ván phụ: Đảo màu quân, Trắng 5p, Đen 4p
    MG->>Players: Thông báo bắt đầu ván phụ phân định
    MG->>TS: Ván phụ kết thúc hòa, xử bên Đen thắng theo quy định
    TS->>DB: Đưa người thắng Trận 2 vào trận Chung kết

    Note over Players,TS: Đếm ngược 30 giây chuyển vòng
    TS->>GW: Vòng 1 hoàn tất, phát đếm ngược nghỉ 30 giây
    GW->>Players: Hiển thị đồng hồ đếm lùi 30 giây chuẩn bị

    Note over Players,DB: Thi đấu Chung kết & Bế mạc
    GW->>TS: Hết 30 giây, tự động khởi tạo Trận Chung kết
    TS->>MG: Tạo phòng đấu Chung kết
    MG->>TS: Trận Chung kết kết thúc, tìm ra Nhà vô địch
    TS->>DB: Cập nhật Nhà vô địch, chuyển status = FINISHED
    DB-->>TS: Lưu thành công
    TS-->>GW: Công bố kết quả chung cuộc
    GW->>Players: Phát sự kiện bế mạc và vinh danh Nhà vô địch"""
    add_code_block("Mermaid - Trình tự Luồng 3: Vòng đời giải đấu loại trực tiếp", mermaid_seq3)

    doc.add_page_break()

    # 2.6 Bảng đặc tả Use Case chi tiết (theo đúng mẫu ex SRS.pdf)
    add_h2("2.6 Bảng đặc tả Use Case chi tiết")

    def render_use_case_tables(uc_meta, main_flow, alt_flow, data_inputs):
        # Bảng 1: Đặc tả Use Case
        tbl_spec = doc.add_table(rows=6, cols=3)
        tbl_spec.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_spec.autofit = False

        # Dòng 1: Mã và Tên
        r0 = tbl_spec.rows[0].cells
        r0[0].paragraphs[0].add_run("Mã Use Case").bold = True
        r0[1].paragraphs[0].add_run(uc_meta["code"]).bold = True
        r0[2].paragraphs[0].add_run(f"Tên Use Case: {uc_meta['name']}").bold = True
        r0[0].width = Inches(1.5)
        r0[1].width = Inches(1.5)
        r0[2].width = Inches(3.3)

        # Dòng 2: Tác nhân
        r1 = tbl_spec.rows[1].cells
        r1[0].paragraphs[0].add_run("Tác nhân").bold = True
        r1[1].paragraphs[0].add_run(uc_meta["actor"])
        r1[0].width = Inches(1.5)
        r1[1].width = Inches(4.8)
        # merge cell 1 và 2 của row 1
        r1[1].merge(r1[2])

        # Dòng 3: Mô tả
        r2 = tbl_spec.rows[2].cells
        r2[0].paragraphs[0].add_run("Mô tả").bold = True
        r2[1].paragraphs[0].add_run(uc_meta["desc"])
        r2[0].width = Inches(1.5)
        r2[1].width = Inches(4.8)
        r2[1].merge(r2[2])

        # Dòng 4: Sự kiện kích hoạt
        r3 = tbl_spec.rows[3].cells
        r3[0].paragraphs[0].add_run("Sự kiện kích hoạt").bold = True
        r3[1].paragraphs[0].add_run(uc_meta["trigger"])
        r3[0].width = Inches(1.5)
        r3[1].width = Inches(4.8)
        r3[1].merge(r3[2])

        # Dòng 5: Tiền điều kiện
        r4 = tbl_spec.rows[4].cells
        r4[0].paragraphs[0].add_run("Tiền điều kiện").bold = True
        r4[1].paragraphs[0].add_run(uc_meta["precond"])
        r4[0].width = Inches(1.5)
        r4[1].width = Inches(4.8)
        r4[1].merge(r4[2])

        # Dòng 6: Hậu điều kiện
        r5 = tbl_spec.rows[5].cells
        r5[0].paragraphs[0].add_run("Hậu điều kiện").bold = True
        r5[1].paragraphs[0].add_run(uc_meta["postcond"])
        r5[0].width = Inches(1.5)
        r5[1].width = Inches(4.8)
        r5[1].merge(r5[2])

        for row in tbl_spec.rows:
            for c in row.cells:
                set_cell_margins(c, 70, 70, 110, 110)
                set_cell_borders(c, "D3D3D3", "D3D3D3", "D3D3D3", "D3D3D3")
                c.paragraphs[0].runs[0].font.name = 'Times New Roman'
                c.paragraphs[0].runs[0].font.size = Pt(11)

        # Đặt nền cho cột thuộc tính
        for idx in range(6):
            set_cell_background(tbl_spec.rows[idx].cells[0], "EBF1F5")

        add_p("", space_before=2, space_after=2)

        # Bảng luồng sự kiện chính
        p_mf = doc.add_paragraph()
        p_mf.add_run("Luồng sự kiện chính (Kịch bản thành công):").bold = True
        p_mf.paragraph_format.space_before = Pt(4)
        p_mf.paragraph_format.space_after = Pt(2)

        tbl_flow = doc.add_table(rows=1, cols=3)
        tbl_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_flow.autofit = False
        hf = tbl_flow.rows[0].cells
        hf[0].paragraphs[0].add_run("STT").bold = True
        hf[1].paragraphs[0].add_run("Thực hiện bởi").bold = True
        hf[2].paragraphs[0].add_run("Hành động chi tiết").bold = True
        for c in hf:
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(c, "2E75B6")
            set_cell_margins(c, 80, 80, 120, 120)
            set_cell_borders(c, "1F4E79", "1F4E79", "1F4E79", "1F4E79")
        hf[0].width = Inches(0.7)
        hf[1].width = Inches(1.6)
        hf[2].width = Inches(4.0)

        for stt, actor, act in main_flow:
            row = tbl_flow.add_row().cells
            row[0].paragraphs[0].add_run(str(stt)).bold = True
            row[1].paragraphs[0].add_run(actor)
            row[2].paragraphs[0].add_run(act)
            for c in row:
                set_cell_background(c, "FFFFFF")
                set_cell_margins(c, 60, 60, 100, 100)
                set_cell_borders(c, "D3D3D3", "D3D3D3", "D3D3D3", "D3D3D3")
                c.paragraphs[0].runs[0].font.name = 'Times New Roman'
                c.paragraphs[0].runs[0].font.size = Pt(11)
            row[0].width = Inches(0.7)
            row[1].width = Inches(1.6)
            row[2].width = Inches(4.0)

        add_p("", space_before=2, space_after=2)

        # Bảng luồng sự kiện thay thế
        if alt_flow:
            p_af = doc.add_paragraph()
            p_af.add_run("Luồng sự kiện thay thế (Ngoại lệ & Rẽ nhánh):").bold = True
            p_af.paragraph_format.space_before = Pt(4)
            p_af.paragraph_format.space_after = Pt(2)

            tbl_alt = doc.add_table(rows=1, cols=3)
            tbl_alt.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl_alt.autofit = False
            ha = tbl_alt.rows[0].cells
            ha[0].paragraphs[0].add_run("STT").bold = True
            ha[1].paragraphs[0].add_run("Thực hiện bởi").bold = True
            ha[2].paragraphs[0].add_run("Hành động xử lý ngoại lệ").bold = True
            for c in ha:
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(c, "C00000") # Dark Red
                set_cell_margins(c, 80, 80, 120, 120)
                set_cell_borders(c, "800000", "800000", "800000", "800000")
            ha[0].width = Inches(0.7)
            ha[1].width = Inches(1.6)
            ha[2].width = Inches(4.0)

            for stt, actor, act in alt_flow:
                row = tbl_alt.add_row().cells
                row[0].paragraphs[0].add_run(str(stt)).bold = True
                row[1].paragraphs[0].add_run(actor)
                row[2].paragraphs[0].add_run(act)
                for c in row:
                    set_cell_background(c, "FFF8F8")
                    set_cell_margins(c, 60, 60, 100, 100)
                    set_cell_borders(c, "E0D0D0", "E0D0D0", "E0D0D0", "E0D0D0")
                    c.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    c.paragraphs[0].runs[0].font.size = Pt(11)
                row[0].width = Inches(0.7)
                row[1].width = Inches(1.6)
                row[2].width = Inches(4.0)

        add_p("", space_before=2, space_after=2)

        # Bảng 2: Dữ liệu đầu vào chi tiết
        if data_inputs:
            p_di = doc.add_paragraph()
            p_di.add_run(f"Dữ liệu đầu vào chức năng ({uc_meta['name']}):").bold = True
            p_di.paragraph_format.space_before = Pt(4)
            p_di.paragraph_format.space_after = Pt(2)

            tbl_data = doc.add_table(rows=1, cols=5)
            tbl_data.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl_data.autofit = False
            hd = tbl_data.rows[0].cells
            hd[0].paragraphs[0].add_run("STT").bold = True
            hd[1].paragraphs[0].add_run("Trường dữ liệu").bold = True
            hd[2].paragraphs[0].add_run("Mô tả").bold = True
            hd[3].paragraphs[0].add_run("Bắt buộc?").bold = True
            hd[4].paragraphs[0].add_run("Điều kiện hợp lệ").bold = True
            for c in hd:
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(c, "1F4E79")
                set_cell_margins(c, 80, 80, 100, 100)
                set_cell_borders(c, "0F2537", "0F2537", "0F2537", "0F2537")
            hd[0].width = Inches(0.5)
            hd[1].width = Inches(1.4)
            hd[2].width = Inches(1.8)
            hd[3].width = Inches(0.9)
            hd[4].width = Inches(1.7)

            for idx, (f_name, desc, req, rule) in enumerate(data_inputs):
                row = tbl_data.add_row().cells
                row[0].paragraphs[0].add_run(str(idx + 1)).bold = True
                row[1].paragraphs[0].add_run(f_name).bold = True
                row[2].paragraphs[0].add_run(desc)
                row[3].paragraphs[0].add_run(req)
                row[4].paragraphs[0].add_run(rule)
                bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
                for c in row:
                    set_cell_background(c, bg)
                    set_cell_margins(c, 60, 60, 90, 90)
                    set_cell_borders(c, "D3D3D3", "D3D3D3", "D3D3D3", "D3D3D3")
                    c.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    c.paragraphs[0].runs[0].font.size = Pt(10.5)
                row[0].width = Inches(0.5)
                row[1].width = Inches(1.4)
                row[2].width = Inches(1.8)
                row[3].width = Inches(0.9)
                row[4].width = Inches(1.7)

        add_p("", space_before=8, space_after=8)

    # -------------------------------------------------------------
    # UC01
    add_h3("2.6.1 Đặc tả Use Case UC01: Đăng ký & Đăng nhập tài khoản")
    uc01_meta = {
        "code": "UC01",
        "name": "Đăng ký & Đăng nhập tài khoản",
        "actor": "Khách, Người chơi có tài khoản",
        "desc": "Cung cấp khả năng đăng ký tài khoản mới và xác thực đăng nhập vào hệ thống nhằm định danh người chơi, cấp phát quyền hạn và bảo vệ phiên làm việc.",
        "trigger": "Người dùng bấm nút Đăng ký hoặc Đăng nhập trên thanh điều hướng đầu trang.",
        "precond": "Người dùng truy cập vào trang web; chưa đăng nhập tài khoản.",
        "postcond": "Người dùng được xác thực thành công, hệ thống cấp phát cặp mã phiên bảo mật và chuyển sang giao diện người chơi chính thức."
    }
    uc01_main = [
        ("1", "Người dùng", "Chọn chức năng Đăng ký hoặc Đăng nhập trên thanh điều hướng."),
        ("2", "Hệ thống", "Hiển thị biểu mẫu nhập liệu tương ứng."),
        ("3", "Người dùng", "Nhập các thông tin cần thiết vào các ô nhập liệu và bấm nút gửi yêu cầu."),
        ("4", "Hệ thống", "Kiểm tra tính đầy đủ của các trường thông tin bắt buộc."),
        ("5", "Hệ thống", "Kiểm tra định dạng địa chỉ thư điện tử và độ dài mật khẩu."),
        ("6", "Hệ thống", "Với Đăng ký: Kiểm tra tính duy nhất của thư điện tử và tên người dùng, băm mật khẩu một chiều và lưu bản ghi người dùng mới với điểm xếp hạng khởi tạo 1200."),
        ("7", "Hệ thống", "Với Đăng nhập: Đối chiếu mật khẩu nhập vào với mật khẩu đã băm trong cơ sở dữ liệu."),
        ("8", "Hệ thống", "Khởi tạo mã truy cập ngắn hạn trong bộ nhớ tạm và mã làm mới dài hạn trong cookie bảo mật chỉ đọc."),
        ("9", "Hệ thống", "Điều hướng người dùng về trang chủ và hiển thị thông tin hồ sơ người chơi.")
    ]
    uc01_alt = [
        ("4a", "Hệ thống", "Báo lỗi: Vui lòng điền đầy đủ các trường thông tin bắt buộc nếu người dùng bỏ trống ô dữ liệu."),
        ("5a", "Hệ thống", "Báo lỗi: Địa chỉ thư điện tử không đúng định dạng quy định."),
        ("5b", "Hệ thống", "Báo lỗi: Mật khẩu phải có độ dài tối thiểu 6 ký tự."),
        ("6a", "Hệ thống", "Báo lỗi: Thư điện tử hoặc tên người dùng đã tồn tại trên hệ thống."),
        ("7a", "Hệ thống", "Báo lỗi: Thông tin tài khoản hoặc mật khẩu không chính xác khi đăng nhập sai thông tin."),
        ("7b", "Hệ thống", "Báo lỗi: Tài khoản của bạn đã bị khóa bởi quản trị viên.")
    ]
    uc01_data = [
        ("username", "Tên định danh người chơi", "Có (khi đăng ký)", "Từ 3 đến 20 ký tự, không chứa ký tự đặc biệt"),
        ("email", "Địa chỉ thư điện tử", "Có", "Đúng định dạng thư điện tử hợp lệ"),
        ("password", "Mật khẩu xác thực", "Có", "Độ dài từ 6 đến 50 ký tự"),
        ("confirmPassword", "Xác nhận lại mật khẩu", "Có (khi đăng ký)", "Khớp với trường password")
    ]
    render_use_case_tables(uc01_meta, uc01_main, uc01_alt, uc01_data)

    # -------------------------------------------------------------
    # UC02
    add_h3("2.6.2 Đặc tả Use Case UC02: Đấu với máy tính")
    uc02_meta = {
        "code": "UC02",
        "name": "Đấu với máy tính",
        "actor": "Khách, Người chơi có tài khoản",
        "desc": "Cho phép người chơi rèn luyện kỹ năng cờ vua bằng cách thi đấu đối kháng với thuật toán máy tính chạy trên luồng ngầm của trình duyệt web.",
        "trigger": "Người dùng bấm chọn mục Đấu với máy trên giao diện menu chính.",
        "precond": "Người dùng đã mở ứng dụng trên trình duyệt web có hỗ trợ JavaScript.",
        "postcond": "Trận đấu được khởi tạo, bàn cờ phản hồi mượt mà theo từng nước đi của người chơi và máy tính."
    }
    uc02_main = [
        ("1", "Người dùng", "Chọn mục Đấu với máy từ menu chính."),
        ("2", "Hệ thống", "Hiển thị hộp thoại tùy chọn: Cấp độ chơi (Dễ, Trung bình, Khó) và Lựa chọn màu quân (Trắng, Đen, Ngẫu nhiên)."),
        ("3", "Người dùng", "Lựa chọn cấp độ mong muốn, chọn màu quân và bấm nút Bắt đầu."),
        ("4", "Hệ thống", "Khởi tạo luồng xử lý ngầm, nạp cấu hình thuật toán tìm kiếm nước đi tương ứng với cấp độ đã chọn."),
        ("5", "Hệ thống", "Khởi tạo bàn cờ với trạng thái ban đầu. Nếu người chơi chọn cầm quân Đen, máy tính tự động thực hiện nước đi đầu tiên cho bên Trắng."),
        ("6", "Người dùng", "Kéo thả hoặc bấm chọn quân cờ để thực hiện nước đi trên bàn cờ."),
        ("7", "Hệ thống", "Kiểm tra nước đi theo luật cờ vua. Nếu hợp lệ, cập nhật vị trí quân cờ và chuyển lượt cho máy tính."),
        ("8", "Hệ thống", "Luồng ngầm tính toán nước phản hồi dựa trên cây tìm kiếm vị trí và gửi kết quả về giao diện hiển thị."),
        ("9", "Hệ thống", "Cập nhật bàn cờ hiển thị nước đi của máy tính và chuyển lại lượt cho người chơi."),
        ("10", "Hệ thống", "Lặp lại các bước 6 đến 9 cho tới khi ván cờ kết thúc bằng kết quả Chiếu hết, Hòa cờ hoặc người chơi bấm nút Đầu hàng.")
    ]
    uc02_alt = [
        ("7a", "Hệ thống", "Từ chối nước đi và trả quân cờ về vị trí cũ nếu nước đi vi phạm luật cờ vua."),
        ("10a", "Người dùng", "Bấm nút Chơi lại để khởi động lại ván cờ mới với tùy chọn ban đầu."),
        ("10b", "Người dùng", "Bấm nút Xin gợi ý, luồng ngầm đề xuất một nước đi có lợi nhất tại thế cờ hiện tại.")
    ]
    uc02_data = [
        ("difficulty", "Cấp độ tính toán của máy", "Có", "Thuộc tập: easy, medium, hard"),
        ("playerSide", "Màu quân người chơi lựa chọn", "Có", "Thuộc tập: white, black, random"),
        ("moveFrom", "Tọa độ ô cờ xuất phát", "Có (khi đi cờ)", "Chuỗi 2 ký tự cột (a-h) và dòng (1-8)"),
        ("moveTo", "Tọa độ ô cờ đích đến", "Có (khi đi cờ)", "Chuỗi 2 ký tự cột (a-h) và dòng (1-8)"),
        ("promotion", "Loại quân chọn khi phong cấp Tốt", "Tùy chọn", "Thuộc tập: q, r, b, n")
    ]
    render_use_case_tables(uc02_meta, uc02_main, uc02_alt, uc02_data)

    # -------------------------------------------------------------
    # UC03
    add_h3("2.6.3 Đặc tả Use Case UC03: Ghép trận đấu xếp hạng")
    uc03_meta = {
        "code": "UC03",
        "name": "Ghép trận đấu xếp hạng",
        "actor": "Người chơi có tài khoản",
        "desc": "Tự động ghép cặp hai kỳ thủ trực tuyến có điểm xếp hạng tương đương, giám sát ván cờ qua giao thức mạng thời gian thực và cập nhật điểm xếp hạng Elo sau khi kết thúc.",
        "trigger": "Người chơi bấm nút Tìm trận xếp hạng tại sảnh chính.",
        "precond": "Người chơi đã đăng nhập tài khoản và không trong trạng thái tham gia ván đấu khác.",
        "postcond": "Ván đấu được tạo lập, kết quả trận đấu và điểm Elo được lưu trữ vào cơ sở dữ liệu."
    }
    uc03_main = [
        ("1", "Người chơi", "Chọn loại hình thời gian thi đấu (Chớp 3 phút, Nhanh 10 phút) và bấm nút Tìm trận xếp hạng."),
        ("2", "Hệ thống", "Đưa người chơi vào hàng chờ ghép trận máy chủ và hiển thị hiệu ứng đang tìm đối thủ."),
        ("3", "Hệ thống", "So khớp người chơi với một kỳ thủ khác trong hàng chờ có mức chênh lệch điểm Elo trong giới hạn cho phép."),
        ("4", "Hệ thống", "Khởi tạo phòng đấu, phân định ngẫu nhiên bên cầm quân Trắng và bên cầm quân Đen, thiết lập đồng hồ thi đấu với mốc thời gian máy chủ."),
        ("5", "Hệ thống", "Phát thông báo ghép trận thành công tới cả hai người chơi kèm thông tin đối thủ."),
        ("6", "Người chơi", "Thực hiện nước đi trong lượt của mình và gửi lên máy chủ qua kết nối mạng thời gian thực."),
        ("7", "Hệ thống", "Kiểm tra định danh kết nối, kiểm tra tính hợp lệ của lượt đi, tính toán và trừ thời gian suy nghĩ vào quỹ giờ bên vừa đi."),
        ("8", "Hệ thống", "Kiểm tra tính hợp lệ của nước đi theo luật cờ vua."),
        ("9", "Hệ thống", "Cập nhật mốc thời gian chuyển lượt và phát tán nước cờ cùng chuỗi FEN mới tới hai người chơi."),
        ("10", "Hệ thống", "Khi ván cờ kết thúc (Chiếu hết, Hết giờ, Hòa cờ, Đầu hàng), máy chủ tính toán điểm Elo biến thiên của hai bên theo công thức quy định."),
        ("11", "Hệ thống", "Lưu bản ghi ván đấu và điểm xếp hạng mới vào cơ sở dữ liệu, đồng thời phát thông báo tổng kết cho hai người chơi.")
    ]
    uc03_alt = [
        ("2a", "Người chơi", "Bấm nút Hủy tìm trận, hệ thống xóa người chơi khỏi hàng chờ và đưa về màn hình chính."),
        ("3a", "Hệ thống", "Mở rộng biên độ chênh lệch điểm Elo sau mỗi khoảng thời gian 10 giây nếu chưa tìm thấy đối thủ ngay lập tức."),
        ("7a", "Hệ thống", "Đồng hồ đếm lùi của bên đang có lượt đi chạm mốc 0, hệ thống xử thua do hết giờ cho bên đó."),
        ("8a", "Hệ thống", "Nước đi không hợp lệ, hệ thống từ chối cập nhật và gửi cảnh báo lỗi về máy khách của người đi."),
        ("9a", "Hệ thống", "Người chơi bị ngắt kết nối mạng, hệ thống kích hoạt cơ chế ân hạn 45 giây để chờ người chơi vào lại.")
    ]
    uc03_data = [
        ("timeControl", "Cấu hình thời gian ván đấu", "Có", "Thuộc tập: bullet_1m, blitz_3m, rapid_10m"),
        ("userId", "Mã định danh người chơi", "Có", "Chuỗi ký tự định danh tài khoản hợp lệ"),
        ("eloRating", "Điểm xếp hạng hiện tại", "Có", "Số nguyên dương lớn hơn 0"),
        ("move", "Chuỗi nước đi dạng đại số", "Có (khi đi cờ)", "Định dạng hợp lệ theo quy định FIDE")
    ]
    render_use_case_tables(uc03_meta, uc03_main, uc03_alt, uc03_data)

    # -------------------------------------------------------------
    # UC04
    add_h3("2.6.4 Đặc tả Use Case UC04: Tạo và tham gia phòng bạn bè")
    uc04_meta = {
        "code": "UC04",
        "name": "Tạo và tham gia phòng bạn bè",
        "actor": "Người chơi có tài khoản",
        "desc": "Cho phép người chơi tạo phòng đấu riêng với mã mời ngẫu nhiên 6 ký tự để mời bạn bè tham gia thi đấu giao hữu không ảnh hưởng tới điểm xếp hạng.",
        "trigger": "Người chơi bấm chọn Tạo phòng bạn bè hoặc Tham gia bằng mã.",
        "precond": "Người chơi đã đăng nhập tài khoản.",
        "postcond": "Hai người chơi vào cùng một phòng đấu riêng biệt và tiến hành ván đấu giao hữu."
    }
    uc04_main = [
        ("1", "Người chơi A", "Chọn chức năng Tạo phòng bạn bè, lựa chọn cấu hình thời gian thi đấu và bấm Tạo phòng."),
        ("2", "Hệ thống", "Khởi tạo phòng đấu riêng, sinh mã phòng gồm 6 ký tự ngẫu nhiên duy nhất và hiển thị lên màn hình."),
        ("3", "Người chơi A", "Sao chép mã phòng và gửi cho Người chơi B."),
        ("4", "Người chơi B", "Chọn chức năng Tham gia phòng, nhập mã phòng gồm 6 ký tự và bấm Vào phòng."),
        ("5", "Hệ thống", "Kiểm tra tính tồn tại của mã phòng và số lượng người hiện có trong phòng."),
        ("6", "Hệ thống", "Đưa Người chơi B vào phòng, thông báo cho Người chơi A và hiển thị trạng thái hai bên đã sẵn sàng."),
        ("7", "Hệ thống", "Bắt đầu ván đấu, hiển thị bàn cờ thời gian thực cho cả hai người chơi."),
        ("8", "Hai người chơi", "Tiến hành thi đấu theo các quy tắc đồng bộ nước đi và kiểm soát thời gian."),
        ("9", "Hệ thống", "Khi ván đấu kết thúc, hiển thị kết quả mà không tính toán biến thiên điểm Elo xếp hạng.")
    ]
    uc04_alt = [
        ("2a", "Người chơi A", "Bấm nút Rời phòng trước khi có người vào, hệ thống hủy phòng đấu và trả về màn hình chính."),
        ("5a", "Hệ thống", "Báo lỗi: Mã phòng không tồn tại hoặc đã hết hạn nếu người chơi nhập sai mã."),
        ("5b", "Hệ thống", "Báo lỗi: Phòng đấu đã đủ 2 người chơi nếu có người thứ ba cố tình tham gia.")
    ]
    uc04_data = [
        ("timeControl", "Thời gian mỗi bên trong phòng", "Có", "Thuộc tập: unlimited, 5m, 10m, 15m"),
        ("roomCode", "Mã tham gia phòng bạn bè", "Có (khi tham gia)", "Chuỗi gồm chính xác 6 ký tự chữ hoa hoặc chữ số")
    ]
    render_use_case_tables(uc04_meta, uc04_main, uc04_alt, uc04_data)

    # -------------------------------------------------------------
    # UC05
    add_h3("2.6.5 Đặc tả Use Case UC05: Tham gia giải đấu loại trực tiếp")
    uc05_meta = {
        "code": "UC05",
        "name": "Tham gia giải đấu loại trực tiếp",
        "actor": "Người chơi có tài khoản",
        "desc": "Người chơi tổ chức hoặc ghi danh tham gia giải đấu thể thức loại trực tiếp 4 hoặc 8 người; hệ thống tự động sinh sơ đồ nhánh đấu, chuyển vòng sau 30 giây nghỉ và tự động kích hoạt ván phụ khi có kết quả hòa cờ.",
        "trigger": "Người chơi bấm chọn Giải đấu từ thanh menu chính.",
        "precond": "Người chơi đã đăng nhập tài khoản.",
        "postcond": "Giải đấu diễn ra tuần tự qua các vòng, tìm ra người chiến thắng chung cuộc và cập nhật danh hiệu."
    }
    uc05_main = [
        ("1", "Chủ phòng", "Chọn tạo giải đấu mới, chọn quy mô (4 hoặc 8 người) và nhận mã mời giải đấu."),
        ("2", "Các kỳ thủ", "Nhập mã mời để ghi danh vào danh sách chờ của giải đấu."),
        ("3", "Chủ phòng", "Khi phòng chờ đã đủ số lượng người đăng ký, bấm nút Bắt đầu giải đấu."),
        ("4", "Hệ thống", "Sinh sơ đồ phân nhánh thi đấu loại trực tiếp ngẫu nhiên và hiển thị cây thi đấu cho toàn thể kỳ thủ."),
        ("5", "Hệ thống", "Tự động tạo các phòng đấu cho các cặp đấu ở Vòng 1 và điều hướng các kỳ thủ vào bàn cờ."),
        ("6", "Kỳ thủ", "Thi đấu ván cờ theo quy định thi đấu trực tuyến thông thường."),
        ("7", "Hệ thống", "Khi ván đấu có người thắng: Cập nhật người thắng tiến vào vòng kế tiếp, người thua dừng bước."),
        ("8", "Hệ thống", "Khi ván đấu hòa: Tự động kích hoạt ván phụ thi đấu nhanh với màu quân đảo chiều, Trắng có 5 phút, Đen có 4 phút và Đen hưởng ưu thế hòa là thắng để phân định dứt điểm người đi tiếp."),
        ("9", "Hệ thống", "Khi toàn bộ các trận ở vòng hiện tại kết thúc: Kích hoạt đồng hồ đếm ngược 30 giây nghỉ chuẩn bị."),
        ("10", "Hệ thống", "Hết 30 giây: Tự động ghép cặp các kỳ thủ thắng cuộc ở vòng trước để tạo các trận đấu ở vòng tiếp theo."),
        ("11", "Hệ thống", "Lặp lại các bước đến khi trận chung kết kết thúc, vinh danh nhà vô địch và hoàn tất giải đấu.")
    ]
    uc05_alt = [
        ("1a", "Chủ phòng", "Rời khỏi phòng trước khi bắt đầu giải, hệ thống hủy giải đấu và gửi thông báo tới các kỳ thủ khác."),
        ("2a", "Hệ thống", "Từ chối yêu cầu tham gia nếu phòng giải đấu đã đủ số lượng người đăng ký."),
        ("4a", "Hệ thống", "Nếu số lượng người tham gia là số lẻ khi bắt đầu, hệ thống tự động chỉ định một kỳ thủ được đi tiếp vào vòng sau mà không cần thi đấu ở vòng đầu.")
    ]
    uc05_data = [
        ("tournamentSize", "Số lượng kỳ thủ trong giải", "Có", "Thuộc tập giá trị: 4, 8"),
        ("tournamentCode", "Mã mời tham gia giải đấu", "Có", "Chuỗi gồm chính xác 6 ký tự chữ hoa"),
        ("tournamentId", "Mã định danh giải đấu", "Hệ thống tự sinh", "Chuỗi định danh duy nhất")
    ]
    render_use_case_tables(uc05_meta, uc05_main, uc05_alt, uc05_data)

    # -------------------------------------------------------------
    # UC06
    add_h3("2.6.6 Đặc tả Use Case UC06: Xem lại và phân tích ván cờ")
    uc06_meta = {
        "code": "UC06",
        "name": "Xem lại và phân tích ván cờ",
        "actor": "Người chơi có tài khoản, Khách",
        "desc": "Cho phép người chơi duyệt lại toàn bộ diễn biến ván cờ đã thi đấu theo từng nước đi, đồng thời sử dụng công cụ phân tích tự động để đánh giá chất lượng nước đi dựa trên độ suy giảm ưu thế.",
        "trigger": "Người chơi bấm nút Xem lại ván cờ sau khi kết thúc trận đấu hoặc chọn một ván đấu từ mục Lịch sử thi đấu.",
        "precond": "Ván đấu đã kết thúc và biên bản nước đi đã được lưu trữ trong cơ sở dữ liệu.",
        "postcond": "Giao diện bàn cờ hiển thị nước cờ kèm đánh giá trực quan về độ chính xác và sai lầm chiến thuật."
    }
    uc06_main = [
        ("1", "Người dùng", "Bấm vào nút Phân tích ván cờ tại màn hình kết thúc ván hoặc chọn từ danh sách lịch sử ván cờ."),
        ("2", "Hệ thống", "Tải toàn bộ danh sách nước đi theo định dạng PGN và thông tin hai bên từ cơ sở dữ liệu."),
        ("3", "Hệ thống", "Hiển thị bàn cờ phân tích với thanh công cụ duyệt nước đi (Nước đầu, Nước trước, Nước sau, Nước cuối)."),
        ("4", "Người dùng", "Bấm nút Phân tích tự động."),
        ("5", "Hệ thống", "Kích hoạt luồng xử lý WebAssembly chạy ngầm công cụ phân tích Stockfish."),
        ("6", "Hệ thống", "Lần lượt gửi chuỗi vị trí FEN của từng nước đi tới công cụ phân tích để tính toán điểm số ưu thế."),
        ("7", "Hệ thống", "Tính toán độ suy giảm ưu thế giữa nước đi thực tế và nước đi tối ưu mà công cụ đề xuất."),
        ("8", "Hệ thống", "Phân loại từng nước cờ thành các cấp bậc: Nước tối ưu, Rất tốt, Tốt, Kém chính xác, Sai lầm, Sai lầm nghiêm trọng."),
        ("9", "Hệ thống", "Vẽ biểu đồ ưu thế của ván cờ theo thời gian và hiển thị mũi tên gợi ý nước đi tối ưu trên bàn cờ."),
        ("10", "Người dùng", "Bấm vào từng nước đi trong danh sách để xem thế cờ và lời giải thích chiến thuật tương ứng.")
    ]
    uc06_alt = [
        ("5a", "Hệ thống", "Trình duyệt không hỗ trợ WebAssembly, hệ thống hiển thị thông báo yêu cầu người dùng cập nhật trình duyệt lên phiên bản mới hơn."),
        ("6a", "Người dùng", "Bấm nút Dừng phân tích, hệ thống tạm dừng luồng tính toán ngầm và giữ nguyên các kết quả đã xử lý được.")
    ]
    uc06_data = [
        ("matchId", "Mã định danh ván đấu cần xem", "Có", "Chuỗi định danh bản ghi ván cờ hợp lệ"),
        ("analysisDepth", "Độ sâu tìm kiếm của công cụ", "Tùy chọn", "Số nguyên dương từ 10 đến 20")
    ]
    render_use_case_tables(uc06_meta, uc06_main, uc06_alt, uc06_data)

    # -------------------------------------------------------------
    # UC07
    add_h3("2.6.7 Đặc tả Use Case UC07: Giải bài tập cờ thế và học cờ")
    uc07_meta = {
        "code": "UC07",
        "name": "Giải bài tập cờ thế và học cờ",
        "actor": "Khách, Người chơi có tài khoản",
        "desc": "Cung cấp kho thế cờ chiến thuật tương tác và các bài giảng bài bản giúp người chơi rèn luyện khả năng phát hiện đòn chiến thuật và nâng cao trình độ.",
        "trigger": "Người dùng bấm chọn mục Bài tập cờ thế hoặc Học cờ trên thanh menu chính.",
        "precond": "Người dùng truy cập vào hệ thống.",
        "postcond": "Hệ thống ghi nhận kết quả giải thế cờ, cập nhật tiến độ hoàn thành bài học vào hồ sơ cá nhân."
    }
    uc07_main = [
        ("1", "Người dùng", "Chọn danh mục Bài tập cờ thế từ menu chính."),
        ("2", "Hệ thống", "Tải và hiển thị bài tập thế cờ phù hợp từ cơ sở dữ liệu kèm yêu cầu chiến thuật (Bên Trắng đi trước và giành ưu thế)."),
        ("3", "Người dùng", "Quan sát bàn cờ và thực hiện nước đi bằng cách kéo thả quân cờ."),
        ("4", "Hệ thống", "Đối chiếu nước đi của người dùng với chuỗi nước đi giải pháp chính xác được lưu trong hệ thống."),
        ("5", "Hệ thống", "Nếu nước đi chính xác: Hệ thống tự động thực hiện nước cờ đáp trả của bên đối kháng theo lời giải bài tập."),
        ("6", "Người dùng", "Tiếp tục đi các nước tiếp theo cho đến khi hoàn thành toàn bộ chuỗi phối hợp chiến thuật."),
        ("7", "Hệ thống", "Hiển thị thông báo chúc mừng hoàn thành xuất sắc, cộng điểm rèn luyện và mở nút chuyển sang bài tập kế tiếp.")
    ]
    uc07_alt = [
        ("4a", "Hệ thống", "Nếu nước đi không đúng lời giải: Hiển thị thông báo Nước cờ chưa tối ưu, vui lòng thử lại và đưa quân cờ về vị trí ban đầu."),
        ("4b", "Người dùng", "Bấm nút Xem lời giải, hệ thống hiển thị chuỗi nước đi mẫu nhưng không tính điểm hoàn thành bài tập cho người dùng.")
    ]
    uc07_data = [
        ("puzzleCategory", "Chủ đề đòn phối hợp chiến thuật", "Tùy chọn", "Thuộc danh mục: Ghim quân, Đòn xiên, Chiếu mở"),
        ("userMove", "Nước đi người dùng thực hiện", "Có", "Tọa độ ô cờ hợp lệ theo luật cờ vua")
    ]
    render_use_case_tables(uc07_meta, uc07_main, uc07_alt, uc07_data)

    doc.add_page_break()

    # 2.7 Bảng ma trận kịch bản kiểm thử
    add_h2("2.7 Bảng ma trận kịch bản kiểm thử")
    add_p("Ma trận kịch bản kiểm thử được thiết lập nhằm xác minh tính đúng đắn của toàn bộ quy tắc nghiệp vụ, luật cờ vua quốc tế và các cơ chế xử lý ngoại lệ đặc thù trong hệ thống:")

    table_tc = doc.add_table(rows=1, cols=6)
    table_tc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tc.autofit = False

    htc = table_tc.rows[0].cells
    htc[0].paragraphs[0].add_run("Mã").bold = True
    htc[1].paragraphs[0].add_run("Tên ca kiểm thử").bold = True
    htc[2].paragraphs[0].add_run("Tiền điều kiện").bold = True
    htc[3].paragraphs[0].add_run("Các bước thực hiện").bold = True
    htc[4].paragraphs[0].add_run("Kết quả mong đợi").bold = True
    htc[5].paragraphs[0].add_run("Đánh giá").bold = True
    for c in htc:
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "1F4E79")
        set_cell_margins(c, 80, 80, 100, 100)
        set_cell_borders(c, "0F2537", "0F2537", "0F2537", "0F2537")

    htc[0].width = Inches(0.6)
    htc[1].width = Inches(1.3)
    htc[2].width = Inches(1.1)
    htc[3].width = Inches(1.3)
    htc[4].width = Inches(1.4)
    htc[5].width = Inches(0.6)

    tc_list = [
        ("TC01", "Kiểm tra hợp lệ nước đi & phong cấp Tốt", "Ván cờ đang thi đấu, Tốt Trắng ở ô e7", "1. Đi Tốt từ e7 lên e8\n2. Chọn phong cấp Hậu", "Nước đi được chấp nhận, ô e8 biến thành quân Hậu Trắng, FEN đồng bộ hai bên", "Đạt"),
        ("TC02", "Ngăn chặn nước đi tự đặt Vua bị chiếu", "Xe đen kiểm soát cột e, Mã trắng che chắn Vua", "1. Người chơi Trắng di chuyển Mã đi nơi khác", "Máy chủ từ chối nước đi, trả quân Mã về vị trí cũ, cảnh báo vi phạm luật cờ", "Đạt"),
        ("TC03", "Thực hiện nước cờ nhập thành hợp lệ", "Vua e1, Xe h1, f1 và g1 trống, chưa từng đi", "1. Kéo quân Vua Trắng từ e1 sang g1", "Vua sang g1, Xe tự động sang f1, quyền nhập thành bị xóa trong chuỗi FEN", "Đạt"),
        ("TC04", "Hết giờ khi đối phương thiếu quân chiếu hết", "Trắng hết giờ, Đen chỉ còn duy nhất 1 Vua", "1. Đồng hồ Trắng về 0\n2. Máy chủ kiểm tra quân", "Trận đấu kết thúc với kết quả Hòa cờ do đối phương không đủ lực lượng chiếu hết", "Đạt"),
        ("TC05", "Khôi phục ván đấu trong 45 giây", "Trận đấu xếp hạng đang diễn ra bình thường", "1. Ngắt mạng bên A\n2. Chờ 15s\n3. Bật mạng, F5", "Người chơi A nhận lại đầy đủ trạng thái ván cờ, đồng hồ tiếp tục chạy, ván đấu tiếp diễn", "Đạt"),
        ("TC06", "Xử thua bỏ cuộc khi quá 45 giây mất mạng", "Người chơi A bị ngắt kết nối mạng", "1. Ngắt kết nối bên A\n2. Chờ quá 45 giây", "Hết 45s, hệ thống tự động xử Người chơi A thua do bỏ cuộc, Người chơi B thắng", "Đạt"),
        ("TC07", "Kích hoạt ván phụ khi hòa giải đấu", "Ván bán kết trong giải kết thúc Hòa cờ", "1. Hai bên hòa theo luật cờ vua", "Hệ thống thông báo ván phụ, tự động đảo màu quân, Trắng 5 phút, Đen 4 phút", "Đạt"),
        ("TC08", "Phân định thắng thua khi ván phụ hòa", "Ván phụ thi đấu nhanh kết thúc hòa cờ", "1. Hết giờ hoặc hòa cờ trong ván phụ", "Hệ thống áp dụng ưu thế hòa: Bên cầm quân Đen được công nhận chiến thắng chung cuộc", "Đạt"),
        ("TC09", "Xử lý hủy giải đấu khi Chủ phòng rời", "Giải đấu ở phòng chờ, đã có 3 người tham gia", "1. Chủ phòng bấm nút thoát giải đấu", "Giải đấu bị hủy bỏ, hệ thống gửi thông báo giải bị hủy tới các kỳ thủ còn lại", "Đạt"),
        ("TC10", "Đồng bộ đồng hồ theo mốc thời gian máy chủ", "Người chơi A suy nghĩ 12 giây rồi đi cờ", "1. Người chơi A đi nước cờ sau 12 giây", "Quỹ giờ bên A bị trừ chính xác 12 giây, mốc thời gian mới được gán cho lượt bên B", "Đạt")
    ]

    for idx, (code, name, pre, steps, expect, res) in enumerate(tc_list):
        row = table_tc.add_row().cells
        row[0].paragraphs[0].add_run(code).bold = True
        row[1].paragraphs[0].add_run(name)
        row[2].paragraphs[0].add_run(pre)
        row[3].paragraphs[0].add_run(steps)
        row[4].paragraphs[0].add_run(expect)
        row[5].paragraphs[0].add_run(res).bold = True
        row[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
        for c in row:
            set_cell_background(c, bg)
            set_cell_margins(c, 60, 60, 80, 80)
            set_cell_borders(c, "D3D3D3", "D3D3D3", "D3D3D3", "D3D3D3")
            c.paragraphs[0].runs[0].font.name = 'Times New Roman'
            c.paragraphs[0].runs[0].font.size = Pt(10)
        row[0].width = Inches(0.6)
        row[1].width = Inches(1.3)
        row[2].width = Inches(1.1)
        row[3].width = Inches(1.3)
        row[4].width = Inches(1.4)
        row[5].width = Inches(0.6)

    doc.add_page_break()

    # ==============================================================================
    # 5. CHƯƠNG 3: YÊU CẦU PHI CHỨC NĂNG
    # ==============================================================================
    add_h1("3. YÊU CẦU PHI CHỨC NĂNG")

    add_h2("3.1 Giao diện người dùng")
    add_p("• Giao diện được thiết kế hiện đại, tinh gọn và trực quan, hỗ trợ hiển thị linh hoạt trên cả màn hình máy tính bàn, máy tính xách tay và thiết bị di động thông minh.")
    add_p("• Bàn cờ hiển thị với độ tương phản cao, hỗ trợ kéo thả quân cờ mượt mà, hiển thị các chấm đánh dấu nước đi hợp lệ khi chạm hoặc chọn quân cờ.")
    add_p("• Âm thanh ván cờ (tiếng di chuyển quân cờ, tiếng ăn quân, tiếng chiếu tướng, tiếng kết thúc trận) được phát tức thì, đồng bộ chính xác với từng hành động trên màn hình.")
    add_p("• Cung cấp thông báo trạng thái rõ ràng, dễ hiểu khi kết nối gặp trục trặc, khi đối thủ tạm ngắt kết nối hoặc khi trận đấu kết thúc.")

    add_h2("3.2 Hiệu năng hệ thống")
    add_p("• Độ trễ truyền tải nước cờ: Trong điều kiện mạng thông thường, thời gian từ khi người chơi gửi nước đi đến khi đối thủ nhận được trên màn hình không vượt quá 100 mili-giây.")
    add_p("• Tối ưu hóa tài nguyên vi xử lý: Áp dụng mô hình đồng hồ thi đấu hướng sự kiện, máy chủ không phát xung liên tục từng giây mà chỉ tính toán khi có nước đi, giúp máy chủ có thể duy trì hàng nghìn phòng đấu đồng thời với mức tải vi xử lý dưới 30%.")
    add_p("• Xử lý tính toán độc lập: Các tác vụ phân tích nước cờ của máy tính và công cụ phân tích được chạy hoàn toàn trên luồng xử lý ngầm của trình duyệt, không gây gián đoạn hoặc đơ cứng giao diện người dùng.")

    add_h2("3.3 Độ tin cậy và tính sẵn sàng")
    add_p("• Khả năng tự phục hồi tiến trình: Hệ thống máy chủ được quản lý bằng trình giám sát tiến trình chuyên dụng, tự động khởi động lại ứng dụng trong vòng dưới 3 giây nếu xảy ra sự cố đột ngột.")
    add_p("• Tính ổn định của ván cờ: Cơ chế ân hạn 45 giây bảo vệ người chơi trước các sự cố mạng chập chờn hoặc thao tác vô tình bấm tải lại trang web, đảm bảo ván cờ không bị gián đoạn oan uổng.")
    add_p("• Thời gian sẵn sàng hoạt động: Hệ thống được cấu hình sẵn sàng phục vụ liên tục với tỷ lệ sẵn sàng đạt trên 99.5% thời gian hoạt động.")

    add_h2("3.4 An toàn và bảo mật dữ liệu")
    add_p("• Mã hóa lưu lượng đường truyền: Toàn bộ dữ liệu trao đổi giữa máy khách và máy chủ qua cả hai giao thức web thông thường và kết nối thời gian thực đều được mã hóa bằng chứng chỉ bảo mật mã hóa giao vận phiên bản mới nhất.")
    add_p("• Bảo mật mật khẩu: Mật khẩu người dùng được băm một chiều bằng thuật toán mã hóa Bcrypt với hệ số muối phù hợp, không bao giờ lưu trữ dưới dạng văn bản thô trong cơ sở dữ liệu.")
    add_p("• Cơ chế xác thực hai lớp mã phiên: Mã truy cập ngắn hạn lưu trong bộ nhớ tạm nhằm ngăn ngừa mã độc đọc lén, mã làm mới dài hạn lưu trong cookie bảo mật chỉ đọc có gắn các cờ bảo vệ chống can thiệp từ mã kịch bản ngoài và chống giả mạo yêu cầu từ trang web khác.")
    add_p("• Kiểm soát tính hợp lệ phía máy chủ: Áp dụng nguyên tắc kiểm tra chặt chẽ, mọi nước cờ, thời gian suy nghĩ và quyền hạn thao tác đều do máy chủ phán quyết, ngăn ngừa hoàn toàn các hành vi gian lận sửa đổi mã nguồn phía máy khách.")

    add_h2("3.5 Khả năng mở rộng")
    add_p("• Cấu trúc phân tầng linh hoạt: Dự án được tổ chức theo mô hình tách biệt rõ ràng giữa giao diện hiển thị và logic nghiệp vụ máy chủ, giúp dễ dàng nâng cấp hoặc thay thế từng phân hệ mà không ảnh hưởng tới toàn bộ hệ thống.")
    add_p("• Khả năng mở rộng máy chủ phân tán: Kiến trúc được thiết kế sẵn sàng để tích hợp lớp cơ sở dữ liệu bộ nhớ đệm trung gian phục vụ chia sẻ trạng thái ván cờ và điều phối tin nhắn giữa nhiều cụm máy chủ khi lượng người dùng tăng cao trong tương lai.")

    return doc

if __name__ == "__main__":
    output_path = r"C:\Document\TTTN\Chess\docs\SRS_Document_Chess_Online.docx"
    print(f"Đang tiến hành tạo tài liệu SRS tại: {output_path}...")
    document = build_srs_document()
    document.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f"Hoàn thành! Kích thước tệp tin Word: {file_size} bytes.")
